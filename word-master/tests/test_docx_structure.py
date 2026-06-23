"""docx_structure.py 的单元测试。

动态构造 docx（python-docx 在 tmp_path 下临时生成），不引入 fixture 二进制文件。
cmd_split/cmd_merge 直接 print 到 stdout，测试通过读产出文件断言。
"""
import sys
import os
from pathlib import Path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))

import docx
import pytest
from docx.enum.style import WD_STYLE_TYPE
from docx_structure import _detect_heading_level


def _make_doc_with_heading(text, level, style_name=None):
    """构造一个含单个标题段落的 docx Document（不落盘）。

    style_name 默认用英文 "Heading N"；传中文样式名时用 add_style 创建。
    """
    doc = docx.Document()
    if style_name is None:
        p = doc.add_heading(text, level=level)
    else:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        p = doc.add_paragraph(text, style=style)
    return doc, p


def test_detect_heading_english():
    """英文 Heading 1 / Heading 3 识别为 level 1 / 3。"""
    doc, p1 = _make_doc_with_heading("标题一", level=1)
    assert _detect_heading_level(p1) == 1
    doc, p3 = _make_doc_with_heading("标题三", level=3)
    assert _detect_heading_level(p3) == 3


def test_detect_heading_chinese():
    """中文 标题 1 / 标题 2 识别为 level 1 / 2。"""
    _doc, p1 = _make_doc_with_heading("中文标题一", level=1, style_name="标题 1")
    assert _detect_heading_level(p1) == 1
    _doc, p2 = _make_doc_with_heading("中文标题二", level=2, style_name="标题 2")
    assert _detect_heading_level(p2) == 2


def test_detect_non_heading():
    """Normal 段落返回 None。"""
    doc = docx.Document()
    p = doc.add_paragraph("正文段落")
    assert _detect_heading_level(p) is None


def test_detect_heading_out_of_range():
    """Heading 7 识别为 level 7（真实级别），但模板只支持 1-6，调用方不编号。"""
    _doc, p7 = _make_doc_with_heading("第七级", level=7)
    assert _detect_heading_level(p7) == 7


from docx.shared import RGBColor
from docx_structure import cmd_numbering_add


def _save_doc(tmp_path, doc, name="test.docx"):
    """保存 docx 到 tmp_path 并返回路径。"""
    p = tmp_path / name
    doc.save(str(p))
    return str(p)


def test_add_basic_chinese_chapter(tmp_path):
    """第{1}章 / {1}.{2} 模板，H1/H2 文本加编号前缀。"""
    doc = docx.Document()
    doc.add_heading("引言", level=1)
    doc.add_heading("背景", level=2)
    doc.add_heading("方法", level=1)
    docx_path = _save_doc(tmp_path, doc)

    cmd_numbering_add(docx_path, {1: "第{1}章 ", 2: "{1}.{2} "}, start_from=1)

    result = docx.Document(docx_path)
    assert result.paragraphs[0].text == "第1章 引言"
    assert result.paragraphs[1].text == "1.1 背景"
    assert result.paragraphs[2].text == "第2章 方法"


def test_add_skips_h7_plus(tmp_path):
    """Heading 7 段落不编号（模板只支持 1-6）。"""
    doc = docx.Document()
    doc.add_heading("H7 标题", level=7)
    doc.add_heading("H1 标题", level=1)
    docx_path = _save_doc(tmp_path, doc)

    cmd_numbering_add(docx_path, {1: "{1} "}, start_from=1)

    result = docx.Document(docx_path)
    assert result.paragraphs[0].text == "H7 标题"
    assert result.paragraphs[1].text == "1 H1 标题"


def test_add_start_from(tmp_path):
    """--start-from 5，第一个 H1 是 第5章。"""
    doc = docx.Document()
    doc.add_heading("引言", level=1)
    docx_path = _save_doc(tmp_path, doc)

    cmd_numbering_add(docx_path, {1: "第{1}章 "}, start_from=5)

    result = docx.Document(docx_path)
    assert result.paragraphs[0].text == "第5章 引言"


def test_add_strips_existing_prefix(tmp_path):
    """标题已有 第1章 前缀，重新 add 不会变成 第1章 第1章。"""
    doc = docx.Document()
    doc.add_heading("第1章 引言", level=1)
    docx_path = _save_doc(tmp_path, doc)

    cmd_numbering_add(docx_path, {1: "第{1}章 "}, start_from=1)

    result = docx.Document(docx_path)
    assert result.paragraphs[0].text == "第1章 引言"


def test_add_multi_run_preserves_format(tmp_path):
    """多 run 标题：编号写进 run[0]，run[1] 文本和格式不变。"""
    doc = docx.Document()
    p = doc.add_heading(level=1)
    run0 = p.add_run("附录")
    run0.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p.add_run("A 测试方法")
    docx_path = _save_doc(tmp_path, doc)

    cmd_numbering_add(docx_path, {1: "第{1}章 "}, start_from=1)

    result = docx.Document(docx_path)
    rp = result.paragraphs[0]
    assert rp.runs[0].text == "第1章 附录"
    assert rp.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert rp.runs[1].text == "A 测试方法"


def test_add_empty_template_skips_level(tmp_path):
    """--h1 为空串时 H1 不编号，H2 正常编号。"""
    doc = docx.Document()
    doc.add_heading("引言", level=1)
    doc.add_heading("背景", level=2)
    docx_path = _save_doc(tmp_path, doc)

    cmd_numbering_add(docx_path, {1: "", 2: "{1}.{2} "}, start_from=1)

    result = docx.Document(docx_path)
    assert result.paragraphs[0].text == "引言"
    assert result.paragraphs[1].text == "0.1 背景"


def test_overwrite_creates_bak(tmp_path):
    """不传 output_path，覆盖原文件并生成 .bak（.bak 内容是改前的）。"""
    doc = docx.Document()
    doc.add_heading("引言", level=1)
    docx_path = _save_doc(tmp_path, doc)
    original_bytes = Path(docx_path).read_bytes()

    cmd_numbering_add(docx_path, {1: "第{1}章 "}, start_from=1)

    bak_path = docx_path + ".bak"
    assert Path(bak_path).exists()
    assert Path(bak_path).read_bytes() == original_bytes
    result = docx.Document(docx_path)
    assert result.paragraphs[0].text == "第1章 引言"


def test_output_flag_no_bak(tmp_path):
    """传 output_path，新文件有编号、原文件未动、无 .bak。"""
    doc = docx.Document()
    doc.add_heading("引言", level=1)
    docx_path = _save_doc(tmp_path, doc)
    original_text = docx.Document(docx_path).paragraphs[0].text
    out_path = str(tmp_path / "out.docx")

    cmd_numbering_add(docx_path, {1: "第{1}章 "}, start_from=1, output_path=out_path)

    assert not Path(docx_path + ".bak").exists()
    assert docx.Document(docx_path).paragraphs[0].text == original_text
    assert docx.Document(out_path).paragraphs[0].text == "第1章 引言"


def test_save_config_writes_yaml(tmp_path):
    """--save-config 写出的 YAML 可被 load_config 读回。"""
    from _numbering_render import load_config as _load

    doc = docx.Document()
    doc.add_heading("引言", level=1)
    docx_path = _save_doc(tmp_path, doc)
    cfg_path = str(tmp_path / "cfg.yaml")

    cmd_numbering_add(
        docx_path, {1: "第{1}章 ", 2: "{1}.{2} "}, start_from=1,
        save_config_path=cfg_path,
    )

    loaded = _load(cfg_path)
    assert loaded[1] == "第{1}章 "
    assert loaded[2] == "{1}.{2} "
    assert loaded["start_from"] == 1