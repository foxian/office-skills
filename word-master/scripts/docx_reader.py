import docx
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _extract_style_props(style):
    """
    提取 Word 样式的具体属性（字体、字号、对齐、间距等）。
    返回一个包含格式信息的字典。
    """
    props = {}
    if not hasattr(style, 'font'):
        return props

    # 提取字体属性
    if style.font.name: props['font'] = style.font.name
    if style.font.size: props['size'] = f"{style.font.size.pt}pt"
    if style.font.bold is not None: props['bold'] = style.font.bold
    if style.font.italic is not None: props['italic'] = style.font.italic

    # 提取段落格式属性
    if hasattr(style, 'paragraph_format'):
        pf = style.paragraph_format
        if pf.alignment is not None:
            align_map = {
                WD_ALIGN_PARAGRAPH.LEFT: 'LEFT',
                WD_ALIGN_PARAGRAPH.CENTER: 'CENTER',
                WD_ALIGN_PARAGRAPH.RIGHT: 'RIGHT',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'JUSTIFY'
            }
            props['align'] = align_map.get(pf.alignment, str(pf.alignment))
        if pf.line_spacing is not None: props['line_spacing'] = pf.line_spacing
        if pf.first_line_indent is not None:
            props['first_line_indent'] = f"{pf.first_line_indent.pt}pt"
        if pf.space_before is not None:
            props['space_before'] = f"{pf.space_before.pt}pt"
        if pf.space_after is not None:
            props['space_after'] = f"{pf.space_after.pt}pt"

    return props

def _extract_run_overrides(paragraph):
    """
    提取段落中各个 Run（文本块）相对于段落样式的手动格式覆盖。
    """
    overrides = []
    for run in paragraph.runs:
        text = run.text.strip()
        if not text:
            continue

        run_props = {}
        # 如果 Run 对象的属性不是 None，说明它被手动显式设置过
        if run.font.name is not None: run_props['font'] = run.font.name
        if run.font.size is not None: run_props['size'] = f"{run.font.size.pt}pt"
        if run.font.bold is not None: run_props['bold'] = run.font.bold
        if run.font.italic is not None: run_props['italic'] = run.font.italic

        if run_props:
            prop_str = ", ".join(f"{k}={v}" for k, v in run_props.items())
            overrides.append(f"{prop_str} on \"{text}\"")
    return overrides

def _extract_cell_overrides(cell):
    """
    提取单元格内所有段落中 Run 的格式覆盖信息列表。
    """
    all_overrides = []
    for para in cell.paragraphs:
        all_overrides.extend(_extract_run_overrides(para))
    return all_overrides

def _extract_table_markdown(table, table_idx: int):
    """
    将单个 python-docx Table 对象渲染为带单元格 ID 标注的 Markdown 行列表。

    格式：
      <!-- TABLE t0: 3 cols × 5 rows -->
      | [t0r0c0] **内容** | ... |
      |---|---|---|
      | [t0r1c0] 内容 | ... |
    """
    rows = table.rows
    cols_count = len(table.columns)
    rows_count = len(rows)
    lines = []
    lines.append(f"<!-- TABLE t{table_idx}: {cols_count} cols × {rows_count} rows -->")

    for r_idx, row in enumerate(rows):
        cells = row.cells
        # 去除合并单元格导致的重复（python-docx 中合并单元格会重复出现）
        seen_tcs = set()
        unique_cells = []
        for cell in cells:
            tc_id = id(cell._tc)
            if tc_id not in seen_tcs:
                seen_tcs.add(tc_id)
                unique_cells.append(cell)

        cell_strs = []
        for c_idx, cell in enumerate(unique_cells):
            # 多段落单元格以 ' / ' 连接
            cell_text = ' / '.join(
                p.text for p in cell.paragraphs if p.text.strip()
            )
            overrides = _extract_cell_overrides(cell)
            cell_id = f"[t{table_idx}r{r_idx}c{c_idx}]"

            if r_idx == 0:
                # 首行视为表头，加粗
                display = f"{cell_id} **{cell_text}**"
            else:
                display = f"{cell_id} {cell_text}"

            if overrides:
                override_str = "; ".join(overrides)
                display += f" {{{override_str}}}"

            cell_strs.append(display)

        lines.append("| " + " | ".join(cell_strs) + " |")

        # 在表头行后插入分隔线
        if r_idx == 0:
            sep_cols = len(unique_cells)
            lines.append("|" + "|".join(["---"] * sep_cols) + "|")

    return lines

def extract_rich_markdown(filepath, overview=False):
    """
    读取 DOCX 文档，按文档阅读顺序交织输出段落和表格的 Markdown 表示。

    :param filepath: DOCX 文件路径
    :param overview: 如果为 True，则只输出标题大纲（Heading），表格不参与输出
    """
    doc = docx.Document(filepath)
    lines = []
    lines.append(f"<!-- DOCUMENT: {filepath} -->")

    if not overview:
        # 收集文档中实际使用到的所有样式名称（仅段落）
        used_style_names = set(p.style.name for p in doc.paragraphs if p.text.strip())
        if used_style_names:
            lines.append("<!-- STYLES:")
            for s_name in sorted(used_style_names):
                if s_name in doc.styles:
                    style = doc.styles[s_name]
                    props = _extract_style_props(style)
                    if props:
                        prop_str = ", ".join(f"{k}:{v}" for k, v in props.items())
                        lines.append(f"  {s_name} -> {prop_str}")
                    else:
                        lines.append(f"  {s_name} -> (default)")
                else:
                    lines.append(f"  {s_name} -> (not found in doc.styles)")
            lines.append("-->")

    lines.append("")  # 头部后的空行

    # 按文档顺序遍历 body 子元素，区分段落（w:p）和表格（w:tbl）
    para_idx = 0
    table_idx = 0

    # 预先建立 python-docx 的段落/表格对象列表，用于通过 XML element 反查
    _para_map = {id(p._element): p for p in doc.paragraphs}
    _table_map = {id(t._element): t for t in doc.tables}

    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            para = _para_map.get(id(child))
            if para is None:
                para_idx += 1
                continue
            text = para.text.strip()
            if not text:
                para_idx += 1
                continue
            style_name = para.style.name if para.style is not None else "Normal"

            if overview and not style_name.startswith('Heading'):
                para_idx += 1
                continue

            overrides = _extract_run_overrides(para) if not overview else []
            if overrides:
                override_str = " | override: " + "; ".join(overrides)
                lines.append(f"### [p{para_idx}] {text} {{{style_name}{override_str}}}")
            else:
                lines.append(f"### [p{para_idx}] {text} {{{style_name}}}")
            para_idx += 1

        elif tag == 'tbl':
            if overview:
                table_idx += 1
                continue
            table = _table_map.get(id(child))
            if table is None:
                table_idx += 1
                continue
            table_lines = _extract_table_markdown(table, table_idx)
            lines.extend(table_lines)
            lines.append("")  # 表格后空行
            table_idx += 1

    return "\n".join(lines)

if __name__ == "__main__":
    import sys
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    filepath = sys.argv[1]
    overview = "--overview" in sys.argv
    print(extract_rich_markdown(filepath, overview))

__all__ = [
    'extract_rich_markdown',
    '_extract_style_props',
    '_extract_run_overrides',
    '_extract_cell_overrides',
    '_extract_table_markdown',
]
