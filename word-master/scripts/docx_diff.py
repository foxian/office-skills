import docx
import re
import sys
import difflib
from dataclasses import dataclass, field
from typing import List, Optional
from docx_reader import extract_rich_markdown

@dataclass
class ParagraphBlock:
    index: int
    text: str
    style: str
    overrides: List[str] = field(default_factory=list)
    raw: str = ""  # Original markdown line

def parse_markdown_blocks(md: str) -> List[ParagraphBlock]:
    """
    Parse rich markdown into structured ParagraphBlock list.
    Extracts: paragraph index [pN], text, style {StyleName}, overrides.
    """
    blocks = []
    # Match: ### [p0] Text {StyleName} or ### [p0] Text {StyleName | override: ...}
    pattern = r'### \[p(\d+)\] (.+?) \{(.+?)\}'
    for line in md.split('\n'):
        match = re.search(pattern, line)
        if match:
            idx = int(match.group(1))
            text = match.group(2).strip()
            style_part = match.group(3)
            # Split style and overrides
            if '|' in style_part:
                style, override_part = style_part.split('|', 1)
                overrides = [o.strip() for o in override_part.replace('override:', '').split(';')]
                overrides = [o for o in overrides if o]
            else:
                style = style_part
                overrides = []
            blocks.append(ParagraphBlock(
                index=idx,
                text=text,
                style=style.strip(),
                overrides=overrides,
                raw=line
            ))
    return blocks

def generate_diff(path1, path2):
    d1 = docx.Document(path1)
    d2 = docx.Document(path2)
    
    texts1 = [p.text for p in d1.paragraphs if p.text.strip()]
    texts2 = [p.text for p in d2.paragraphs if p.text.strip()]
    
    diff = difflib.unified_diff(texts1, texts2, lineterm='', n=0)
    return '\n'.join(diff)


def compare_tables(doc1, doc2) -> str:
    """
    对比两个文档的表格内容（按单元格文本）。
    返回 Markdown 格式的 '## Tables' 章节字符串；无表格时返回空字符串。
    """
    tables1 = doc1.tables
    tables2 = doc2.tables
    max_tables = max(len(tables1), len(tables2))

    if max_tables == 0:
        return ""

    lines = ["## Tables", ""]

    for t_idx in range(max_tables):
        if t_idx >= len(tables1):
            # 文档2 新增的表格
            t2 = tables2[t_idx]
            lines.append(f"### t{t_idx} [新增, {len(t2.columns)}×{len(t2.rows)}]")
            lines.append("+ 整个表格为新增")
            lines.append("")
            continue
        if t_idx >= len(tables2):
            # 文档1 有但文档2 删除的表格
            t1 = tables1[t_idx]
            lines.append(f"### t{t_idx} [删除, {len(t1.columns)}×{len(t1.rows)}]")
            lines.append("- 整个表格被删除")
            lines.append("")
            continue

        t1 = tables1[t_idx]
        t2 = tables2[t_idx]
        rows1 = t1.rows
        rows2 = t2.rows
        max_rows = max(len(rows1), len(rows2))
        cell_diffs = []

        for r_idx in range(max_rows):
            if r_idx >= len(rows1) or r_idx >= len(rows2):
                cell_diffs.append(
                    f"- 行数不同：文档1 {len(rows1)} 行，文档2 {len(rows2)} 行"
                )
                break
            cells1 = rows1[r_idx].cells
            cells2 = rows2[r_idx].cells
            max_cols = max(len(cells1), len(cells2))
            for c_idx in range(max_cols):
                if c_idx >= len(cells1) or c_idx >= len(cells2):
                    cell_diffs.append(
                        f"- [t{t_idx}r{r_idx}] 列数不同：文档1 {len(cells1)} 列，文档2 {len(cells2)} 列"
                    )
                    break
                text1 = cells1[c_idx].text
                text2 = cells2[c_idx].text
                if text1 != text2:
                    cell_diffs.append(
                        f'- [t{t_idx}r{r_idx}c{c_idx}] "{text1}" → "{text2}"'
                    )

        lines.append(
            f"### t{t_idx} ({len(t1.columns)}×{len(t1.rows)})"
        )
        if cell_diffs:
            lines.extend(cell_diffs)
        else:
            lines.append("（无变更）")
        lines.append("")

    return "\n".join(lines)

def align_paragraphs(blocks1: List[ParagraphBlock], blocks2: List[ParagraphBlock]) -> List[tuple]:
    """
    Align paragraphs between two documents.
    Returns: List[(block1_or_None, block2_or_None, status)]
    status: 'same', 'modified', 'new', 'deleted'
    """
    aligned = []

    # Build a map from original index to block for blocks2
    idx2_map = {b.index: b for b in blocks2}

    for i in range(len(blocks1)):
        b1 = blocks1[i]
        b2 = idx2_map.get(b1.index)

        if b2 is None:
            # Original index not in blocks2 -> deleted
            aligned.append((b1, None, 'deleted'))
        elif b1.text == b2.text and b1.style == b2.style and b1.overrides == b2.overrides:
            # Same
            aligned.append((b1, b2, 'same'))
        else:
            # Modified
            aligned.append((b1, b2, 'modified'))

    # Remaining paragraphs in blocks2 are new (fixed version)
    aligned_indices = {b1.index for b1 in blocks1}
    for b2 in blocks2:
        if b2.index not in aligned_indices:
            aligned.append((None, b2, 'new'))

    return aligned

@dataclass
class DiffResult:
    type: str  # 'same', 'content', 'format', 'new', 'deleted'
    original: Optional[ParagraphBlock]
    modified: Optional[ParagraphBlock]
    fields: List[dict] = field(default_factory=list)

def compute_field_diff(block1: Optional[ParagraphBlock], block2: Optional[ParagraphBlock]) -> DiffResult:
    """Compute field-level differences between two paragraph blocks."""
    if block1 is None and block2 is not None:
        return DiffResult(type='new', original=None, modified=block2, fields=[])
    if block1 is not None and block2 is None:
        return DiffResult(type='deleted', original=block1, modified=None, fields=[])
    if block1 is None and block2 is None:
        return DiffResult(type='same', original=None, modified=None, fields=[])

    fields = []
    change_type = 'same'

    # Style change
    if block1.style != block2.style:
        fields.append({'name': 'style', 'orig': block1.style, 'modif': block2.style})
        change_type = 'format'

    # Overrides change
    if block1.overrides != block2.overrides:
        fields.append({'name': 'overrides', 'orig': block1.overrides, 'modif': block2.overrides})
        change_type = 'format'

    # Content change
    if block1.text != block2.text:
        fields.append({'name': 'content', 'orig': block1.text, 'modif': block2.text})
        change_type = 'content'

    return DiffResult(type=change_type, original=block1, modified=block2, fields=fields)

def format_diff_report(diff_results: List[DiffResult], original: str, modified: str, context: int = 1, include_same: bool = False) -> str:
    """Format diff results into Markdown report."""
    lines = []
    lines.append(f"# DOCX Diff Report")
    lines.append(f"<!-- ORIGINAL: {original} -->")
    lines.append(f"<!-- MODIFIED: {modified} -->")
    lines.append("")
    lines.append("## Paragraphs")
    lines.append("")

    for result in diff_results:
        if result.type == 'same':
            if include_same:
                lines.append(f"### [p{result.original.index}] {{{result.original.style}}}")
                lines.append(result.original.text)
                lines.append("")
            continue
        elif result.type == 'new':
            lines.append(f"### [新增] [p{result.modified.index}] {{{result.modified.style}}}")
            lines.append(result.modified.text)
            lines.append("")
        elif result.type == 'deleted':
            lines.append(f"### [删除] [p{result.original.index}] {{{result.original.style}}}")
            lines.append(f"~~{result.original.text}~~")
            lines.append("")
        else:
            # modified
            lines.append(f"### [p{result.original.index}] {{{result.modified.style if result.modified else result.original.style}}}")
            if result.fields:
                lines.append("#### Fields")
                lines.append("| Field | Original | Modified |")
                lines.append("|-------|----------|----------|")
                for f in result.fields:
                    lines.append(f"| {f['name']} | {f.get('orig', '')} | {f.get('modif', '')} |")
            lines.append("")

    return "\n".join(lines)

def generate_rich_diff(path1: str, path2: str, context: int = 1) -> str:
    """
    Generate rich-text semantic diff between two DOCX files.
    """
    import os
    if not os.path.exists(path1):
        raise FileNotFoundError(f"File not found: {path1}")
    if not os.path.exists(path2):
        raise FileNotFoundError(f"File not found: {path2}")

    try:
        md1 = extract_rich_markdown(path1)
        md2 = extract_rich_markdown(path2)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX files: {e}")

    # Parse into blocks
    blocks1 = parse_markdown_blocks(md1)
    blocks2 = parse_markdown_blocks(md2)

    # Align paragraphs
    aligned = align_paragraphs(blocks1, blocks2)

    # Compute diffs
    diff_results = []
    for b1, b2, status in aligned:
        if status == 'same':
            diff_results.append(DiffResult(type='same', original=b1, modified=b2, fields=[]))
        else:
            diff_results.append(compute_field_diff(b1, b2))

    # Format paragraph report
    para_report = format_diff_report(
        diff_results, original=path1, modified=path2,
        context=context, include_same=context > 0
    )

    # Format table report
    d1 = docx.Document(path1)
    d2 = docx.Document(path2)
    table_report = compare_tables(d1, d2)

    if table_report:
        return para_report + "\n\n" + table_report
    return para_report


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description='DOCX Rich Diff')
    parser.add_argument('path1', help='Original DOCX file')
    parser.add_argument('path2', help='Modified DOCX file')
    parser.add_argument('--output', '-o', help='Output file path')
    parser.add_argument('--verbose', '-v', action='store_true', help='Also output to stdout')
    parser.add_argument('--context', '-c', type=int, default=1, help='Context paragraphs count')
    args = parser.parse_args()

    try:
        result = generate_rich_diff(args.path1, args.path2, context=args.context)
    except (FileNotFoundError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(result)
        if args.verbose or not sys.stdout.isatty():
            print(result)
    else:
        print(result)
