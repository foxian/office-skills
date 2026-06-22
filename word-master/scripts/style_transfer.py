import docx
import json
import sys
import os

# Reuse fingerprint computation from style_analyzer
sys.path.insert(0, os.path.dirname(__file__))
from style_analyzer import _detect_list_type, _is_western_font
from validate_style_profile import validate_profile
from style_analyzer import extract_fingerprints

_FALLBACK_MAP = {
    "Body Text": "Normal",
    "Compact": "List Paragraph",
}

_BODY_STYLES_PRIORITY = ["Normal Indent", "Normal", "Body Text", "First Paragraph"]

def _is_body_style(name):
    name_lower = name.lower()
    return (
        name in _BODY_STYLES_PRIORITY 
        or "body" in name_lower 
        or name_lower == "normal"
    )

def _find_fallback_body_style(available_roles):
    for style in _BODY_STYLES_PRIORITY:
        if style in available_roles:
            return style
    return None



def _score(fp, role_fp):
    """
    Calculate similarity score between two fingerprints (0.0 ~ 1.0).
    Weights: size=0.25, bold=0.125, italic=0.125, align=0.125,
             font=0.125, color=0.125, space_before=0.0625, space_after=0.0625.
    line_spacing is recorded in fingerprint but not scored.
    Size uses linear decay: max(0, 1 - diff/10).

    When a field in fp is None (not set, relies on style default), that
    dimension is skipped and the score is renormalized so the max is 1.0.
    This prevents "fake style" documents (no run-level overrides) from
    being penalized for missing font/size information.
    """
    score = 0.0
    weight_sum = 0.0  # tracks total weight of scored dimensions

    # size (weight=0.25): linear decay
    s1 = fp.get("size")
    s2 = role_fp.get("size")
    if s1 is not None:
        weight_sum += 0.25
        if s1 and s2:
            try:
                diff = abs(float(s1.rstrip("pt")) - float(s2.rstrip("pt")))
                score += max(0.0, 1.0 - diff / 10.0) * 0.25
            except ValueError:
                pass
        elif s1 == s2:  # both None or both empty
            score += 0.25

    # bold (weight=0.125)
    if fp.get("bold") is not None:
        weight_sum += 0.125
        score += 0.125 if fp.get("bold") == role_fp.get("bold") else 0.0

    # italic (weight=0.125)
    if fp.get("italic") is not None:
        weight_sum += 0.125
        score += 0.125 if fp.get("italic") == role_fp.get("italic") else 0.0

    # align (weight=0.125)
    if fp.get("align") is not None:
        weight_sum += 0.125
        score += 0.125 if fp.get("align") == role_fp.get("align") else 0.0

    # font (weight=0.125): exact match
    f1 = fp.get("font")
    f2 = role_fp.get("font")
    if f1 is not None:
        weight_sum += 0.125
        if f2 is None:
            pass  # fp has font, role doesn't — no score
        elif f1 == f2:
            score += 0.125

    # color (weight=0.125): exact match on combined "rgb:XXX" or "theme:XXX" string
    c1 = fp.get("color")
    c2 = role_fp.get("color")
    if c1 is not None:
        weight_sum += 0.125
        if c2 is not None and c1 == c2:
            score += 0.125

    # space_before (weight=0.0625): numeric comparison within 0.5pt tolerance
    sb1 = fp.get("space_before")
    if sb1 is not None:
        weight_sum += 0.0625
        sb2 = role_fp.get("space_before") or "0.0pt"
        try:
            if abs(float(sb1.rstrip("pt")) - float(sb2.rstrip("pt"))) < 0.5:
                score += 0.0625
        except (ValueError, AttributeError):
            pass

    # space_after (weight=0.0625): numeric comparison within 0.5pt tolerance
    sa1 = fp.get("space_after")
    if sa1 is not None:
        weight_sum += 0.0625
        sa2 = role_fp.get("space_after") or "0.0pt"
        try:
            if abs(float(sa1.rstrip("pt")) - float(sa2.rstrip("pt"))) < 0.5:
                score += 0.0625
        except (ValueError, AttributeError):
            pass

    # Renormalize: if some dimensions were skipped, scale score to 0-1 range
    if weight_sum == 0:
        return 0.0
    # Penalize when too few dimensions contributed: if less than half the max
    # weight (1.0) was available, scale down proportionally. This prevents
    # paragraphs with all-None fingerprints from scoring 1.0 on just 1-2 dims.
    confidence = min(1.0, weight_sum / 0.5)  # need at least 0.5 weight to be fully trusted
    return (score / weight_sum) * confidence


def match_fingerprint_to_role(fp, profile, threshold=0.6):
    """
    Match a paragraph fingerprint to the nearest role in style_profile.
    Returns best-matching Word style name, or None if best score < threshold.
    """
    best_role = None
    best_score = 0.0
    for entry in profile.get("roles", []):
        s = _score(fp, entry["fingerprint"])
        if s > best_score:
            best_score = s
            best_role = entry["role"]
    return best_role if best_score >= threshold else None


def generate_apply_ops(draft_path, profile, skip_head=0, skip_tail=0):
    """
    Walk draft.docx paragraphs, match roles via deterministic cascade pipeline.
    Pipeline:
      1. Level 0: Direct Heading Name Match
      2. Level 1: XML Numbering (numPr) & Text Bullet prefix Match (Strict Handoff)
      3. Level 2: Body-style (Normal/Body Text) — text pattern first, then default Normal
      4. Level 3: Absolute Fallback to 'Normal'
    """
    # --- Phase 1: generate update_style_definition ops ---
    ops = []
    for entry in profile.get("roles", []):
        role = entry.get("role")
        fp = entry.get("fingerprint")
        if role and fp:
            ops.append({
                "op": "update_style_definition",
                "style": role,
                "fingerprint": fp,
            })


    # --- Phase 2: generate apply_style ops per paragraph ---
    doc = docx.Document(draft_path)
    paras = doc.paragraphs
    n = len(paras)
    available_roles = {entry["role"] for entry in profile.get("roles", [])}
    _normal_fp = next((e["fingerprint"] for e in profile.get("roles", []) if e["role"] == "Normal"), None)

    for i, para in enumerate(paras):
        if i < skip_head or i >= n - skip_tail:
            continue
        style_name = para.style.name if para.style is not None else "Normal"
        role = None

        # Level 0: Direct Heading Match
        if style_name.startswith("Heading ") and style_name in available_roles:
            role = style_name

        # Level 1: XML List / Bullet Prefix Match with Strict Handoff
        if role is None:
            list_type = _detect_list_type(para, doc)
            if list_type is not None:
                if list_type in available_roles:
                    role = list_type
                elif "List Paragraph" in available_roles:
                    role = "List Paragraph"
                else:
                    role = "Normal"

        # Level 1.5: List Paragraph passthrough
        # If the paragraph uses a list-like style (e.g. "List Paragraph", "List Bullet")
        # but Level 1 didn't catch it (no numPr), preserve the original style so that
        # list indentation and formatting are not clobbered by Normal fallback.
        if role is None and ("list" in style_name.lower()) and style_name in doc.styles:
            role = style_name

        # Level 2: 精确匹配（仅当 role 未被设置时）
        if role is None:
            if style_name in available_roles:
                role = style_name
            elif style_name in _FALLBACK_MAP:
                fallback = _FALLBACK_MAP[style_name]
                if fallback in available_roles:
                    role = fallback

        # Level 2.5: 智能正文大类退级
        if role is None and _is_body_style(style_name):
            role = _find_fallback_body_style(available_roles)

        # Level 3: Absolute Fallback
        if role is None:
            role = "Normal" if "Normal" in available_roles else None

        if role:
            is_style_changing = (style_name != role)
            ops.append({
                "op": "apply_style",
                "target": f"p{i}",
                "style": role,
                "clear_run_formats": is_style_changing,
            })


    # --- Phase 3: generate apply_table_style ops per table ---
    table_roles = profile.get("table_roles", [])
    if not table_roles:
        print("[INFO] table_roles 为空或不存在，跳过表格样式迁移。")
    else:
        from table_analyzer import get_table_cols, detect_has_header_row
        for i, table in enumerate(doc.tables):
            cols = get_table_cols(table)
            rows = len(table.rows)
            has_header = detect_has_header_row(table)

            best_role = None
            best_score = -1.0

            for entry in table_roles:
                struct = entry.get("structure", {})
                e_cols = struct.get("cols")
                e_rows = struct.get("rows")
                e_has_header = struct.get("has_header_row")

                score = 0.0
                
                # 1. cols 相同 → +0.5
                if cols == e_cols:
                    score += 0.5
                
                # 2. has_header_row 一致 → +0.3
                if has_header == e_has_header:
                    score += 0.3
                
                # 3. 行数差值 ≤ 2 → +0.2
                if e_rows is not None:
                    try:
                        diff = abs(rows - e_rows)
                        if diff <= 2:
                            score += 0.2
                    except (ValueError, TypeError):
                        pass

                if score > best_score:
                    best_score = score
                    best_role = entry

            if best_role is not None:
                # 智能注入主正文的中文字体作为表格文字的默认 east_asia 字体，防止表格内中文变成宋体或系统默认
                # 我们先在 profile 的 roles 中寻找主正文字体（即第一个非西文字体）
                main_chinese_font = None
                for entry in profile.get("roles", []):
                    f = entry.get("fingerprint", {}).get("font")
                    if f and not _is_western_font(f):
                        main_chinese_font = f
                        break
                
                # 复制一份 best_role 避免直接修改 profile 并展开标准化纠偏与智能设置
                role_copy = json.loads(json.dumps(best_role))
                
                # 1. 规范表头格式：优先保留源模板表头样式，缺失则智能回退 (黑体三号, 加粗, 居中)
                orig_header = best_role.get("header_text") or {}
                header_font = orig_header.get("font")
                header_chinese = (
                    orig_header.get("east_asia")
                    or (header_font if header_font and not _is_western_font(header_font) else None)
                    or "黑体"
                )
                role_copy["header_text"] = {
                    "font": header_font or "黑体",
                    "east_asia": header_chinese,
                    "size": orig_header.get("size") or "16.0pt",  # 三号字
                    "bold": orig_header.get("bold") if orig_header.get("bold") is not None else True,
                    "align": orig_header.get("align") or "center"
                }
                
                # 2. 规范表格内容格式：优先保留源模板正文样式，缺失则智能回退 (仿宋四号, 不加粗, 自适应对齐)
                orig_body = best_role.get("body_text") or {}
                body_font = orig_body.get("font")
                body_chinese = (
                    orig_body.get("east_asia")
                    or (body_font if body_font and not _is_western_font(body_font) else None)
                    or main_chinese_font
                    or "仿宋_GB2312"
                )
                body_size = orig_body.get("size") or "14.0pt"  # 四号字
                body_align = orig_body.get("align") or "justify"
                
                role_copy["body_text"] = {
                    "font": body_font or "Arial",   # 西文
                    "east_asia": body_chinese,
                    "size": body_size,
                    "bold": orig_body.get("bold") if orig_body.get("bold") is not None else False,
                    "align": body_align
                }

                ops.append({
                    "op": "apply_table_style",
                    "target": f"t{i}",
                    "table_role": role_copy
                })

    return ops


def run(template_path, draft_path, output_path,
        profile_path=None, review=False, skip_head=0, skip_tail=0):
    """
    Main entry: two-phase style transfer workflow.
    profile_path: if provided, skip analysis and LLM inference, reuse existing profile.
    review: if True, pause after profile load for user confirmation.
    """
    # --- File existence validation ---
    if not os.path.exists(draft_path):
        raise FileNotFoundError(f"[ERROR] Draft file not found: {draft_path}")
    if profile_path and not os.path.exists(profile_path):
        raise FileNotFoundError(f"[ERROR] Profile file not found: {profile_path}")
    if not profile_path and template_path and not os.path.exists(template_path):
        raise FileNotFoundError(f"[ERROR] Template file not found: {template_path}")

    # --- Phase 1: Get style_profile ---

    # --- Validate profile format ---

    if profile_path:
        print(f"[INFO] Loading existing profile: {profile_path}")
        with open(profile_path, "r", encoding="utf-8") as f:
            profile = json.load(f)

        if not validate_profile(profile):
            raise ValueError(f"[ERROR] style_profile.json 格式验证失败，请检查上述错误")
    else:
        
        print(f"[INFO] Analyzing template: {template_path}")
        fingerprints = extract_fingerprints(template_path)

        # --- LLM inference placeholder ---
        # Current version saves fingerprints to JSON for external LLM to generate style_profile.
        # TODO: integrate API call to auto-generate style_profile.
        fingerprints_path = "fingerprints.json"
        with open(fingerprints_path, "w", encoding="utf-8") as f:
            json.dump(fingerprints, f, ensure_ascii=False, indent=2)
        print(f"[INFO] Fingerprints saved to {fingerprints_path}")
        print("[INFO] Please review fingerprints.json, ask LLM to generate style_profile.json,")
        print("       then re-run with: --profile style_profile.json")
        return

    # --- Review mode: pause for user confirmation ---
    if review:
        print("[REVIEW] Profile loaded. Edit style_profile.json if needed, then press Enter to continue...")
        input()

    # --- Phase 2: Generate DSL and delegate to docx_editor ---
    print(f"[INFO] Generating apply_style operations for: {draft_path}")
    ops = generate_apply_ops(draft_path, profile,
                              skip_head=skip_head, skip_tail=skip_tail)
    print(f"[INFO] Generated {len(ops)} operations")

    # Delegate to docx_editor (reuses backup and safe execution)
    editor_path = os.path.join(os.path.dirname(__file__), "docx_editor.py")
    from docx_editor import apply_operations
    apply_operations(draft_path, ops, output_path)
    print(f"[INFO] Done → {output_path}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Word document style transfer tool.")
    parser.add_argument("template", nargs="?", help="Source template DOCX (skip if using --profile)")
    parser.add_argument("draft", help="Target draft DOCX to apply styles to")
    parser.add_argument("output", help="Output DOCX path")
    parser.add_argument("--profile", help="Reuse existing style_profile.json (skip analysis)")
    parser.add_argument("--review", action="store_true", help="Pause after profile load for manual review")
    parser.add_argument("--skip-head", type=int, default=0, metavar="N", help="Skip first N paragraphs")
    parser.add_argument("--skip-tail", type=int, default=0, metavar="N", help="Skip last N paragraphs")
    args = parser.parse_args()
    run(
        template_path=args.template,
        draft_path=args.draft,
        output_path=args.output,
        profile_path=args.profile,
        review=args.review,
        skip_head=args.skip_head,
        skip_tail=args.skip_tail,
    )
