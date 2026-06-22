import docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys
import os
import json


def _load_template(template_arg):
    """Load format rules from template.

    Args:
        template_arg: Either "json:xxx" (load format_rules/xxx.json) or
                     "docx:xxx" (load styles from xxx.docx)

    Returns:
        dict: Format rules for heading1, heading2, body
    """
    if template_arg.startswith("json:"):
        template_name = template_arg.split(":", 1)[1]
        skill_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_path = os.path.join(skill_dir, "format_rules", f"{template_name}.json")
        with open(template_path, "r", encoding="utf-8") as f:
            return json.load(f)
    elif template_arg.startswith("docx:"):
        docx_path = template_arg.split(":", 1)[1]
        return _load_styles_from_docx(docx_path)
    else:
        return None


def _load_styles_from_docx(docx_path):
    """Extract style rules from a DOCX file."""
    doc = docx.Document(docx_path)
    styles = {}
    for style in doc.styles:
        if style.type == 1:
            name = style.name.lower()
            if "heading 1" in name or "标题1" in name:
                styles["heading1"] = _style_to_dict(style)
            elif "heading 2" in name or "标题2" in name:
                styles["heading2"] = _style_to_dict(style)
            elif "normal" in name or "正文" in name:
                styles["body"] = _style_to_dict(style)
    if not styles:
        styles = {"heading1": {}, "heading2": {}, "body": {}}
    return styles


def _style_to_dict(style):
    """Convert a paragraph style to a dict format."""
    result = {}
    try:
        if style.font.name:
            result["font"] = style.font.name
        if style.font.size:
            result["size"] = f"{style.font.size.pt}pt"
        if style.font.bold is not None:
            result["bold"] = style.font.bold
        if style.alignment:
            align_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
            }
            result["alignment"] = align_map.get(style.alignment, "left")
        if style.paragraph_format.line_spacing:
            result["line_spacing"] = style.paragraph_format.line_spacing
        if style.paragraph_format.first_line_indent:
            result["first_line_indent"] = f"{style.paragraph_format.first_line_indent.pt}pt"
    except:
        pass
    return result


def _set_run_east_asia(run, font_name):
    """Set East Asian font on a run via rFonts/@w:eastAsia."""
    rpr = run._element.find(qn('w:rPr'))
    if rpr is None:
        rpr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)


def _apply_format_to_paragraph(paragraph, fmt):
    """Apply format dict to a paragraph."""
    if not fmt:
        return
    pf = paragraph.paragraph_format
    if "font" in fmt:
        for run in paragraph.runs:
            run.font.name = fmt["font"]
    if "east_asia" in fmt:
        for run in paragraph.runs:
            _set_run_east_asia(run, fmt["east_asia"])
    if "size" in fmt:
        size_pt = float(fmt["size"].rstrip("pt"))
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)
    if "bold" in fmt:
        for run in paragraph.runs:
            run.bold = fmt["bold"]
    if "alignment" in fmt:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        paragraph.alignment = align_map.get(fmt["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
    if "line_spacing" in fmt:
        pf.line_spacing = fmt["line_spacing"]
    if "first_line_indent" in fmt:
        indent_str = fmt["first_line_indent"]
        if indent_str.endswith("em"):
            pf.first_line_indent = Emu(int(float(indent_str.rstrip("em")) * 914400 / 2))
        elif indent_str.endswith("pt"):
            pf.first_line_indent = Pt(float(indent_str.rstrip("pt")))
    if "space_before" in fmt:
        pf.space_before = Pt(float(fmt["space_before"].rstrip("pt")))
    if "space_after" in fmt:
        pf.space_after = Pt(float(fmt["space_after"].rstrip("pt")))


def convert_markdown(md_path, out_path, template=None):
    doc = docx.Document()
    format_rules = None

    if template:
        format_rules = _load_template(template)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            if format_rules and "heading1" in format_rules:
                _apply_format_to_paragraph(p, format_rules["heading1"])
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            if format_rules and "heading2" in format_rules:
                _apply_format_to_paragraph(p, format_rules["heading2"])
        else:
            p = doc.add_paragraph(line)
            if format_rules and "body" in format_rules:
                _apply_format_to_paragraph(p, format_rules["body"])

    doc.save(out_path)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: md_to_docx.py <input.md> <output.docx> [--template <template>]")
        sys.exit(1)

    md_path = sys.argv[1]
    out_path = sys.argv[2]
    template = None

    if "--template" in sys.argv:
        idx = sys.argv.index("--template")
        template = sys.argv[idx + 1]

    convert_markdown(md_path, out_path, template)