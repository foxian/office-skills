import docx
import json
import re
import shutil
import sys
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _replace_text_in_runs(paragraph, find_txt, repl_txt):
    """Run-level text replacement to preserve inline formatting (bold, italic, font, etc.)."""
    full_text = paragraph.text
    if find_txt not in full_text:
        return False

    replaced = False
    for run in paragraph.runs:
        if find_txt in run.text:
            run.text = run.text.replace(find_txt, repl_txt)
            replaced = True

    if not replaced:
        paragraph.text = full_text.replace(find_txt, repl_txt)
        replaced = True

    return replaced

def _copy_run_format(src_run, dest_run):
    """Copy formatting from src_run to dest_run"""
    dest_run.bold = src_run.bold
    dest_run.italic = src_run.italic
    dest_run.underline = src_run.underline
    dest_run.style = src_run.style
    if src_run.font:
        if src_run.font.name:
            dest_run.font.name = src_run.font.name
        if src_run.font.size:
            dest_run.font.size = src_run.font.size
        if src_run.font.color and src_run.font.color.rgb:
            dest_run.font.color.rgb = src_run.font.color.rgb
        if src_run.font.highlight_color:
            dest_run.font.highlight_color = src_run.font.highlight_color
        dest_run.font.strike = src_run.font.strike
        dest_run.font.subscript = src_run.font.subscript
        dest_run.font.superscript = src_run.font.superscript

def _apply_font_attributes(run, params):
    if "name" in params:
        run.font.name = params["name"]
    if "east_asia" in params:
        from docx.oxml.ns import qn
        rpr = run._element.find(qn('w:rPr'))
        if rpr is None:
            rpr = run._element.makeelement(qn('w:rPr'), {})
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:eastAsia'), params["east_asia"])
    if "size" in params:
        size_str = str(params["size"])
        if size_str.endswith("pt"):
            from docx.shared import Pt
            run.font.size = Pt(float(size_str.rstrip("pt")))
        elif size_str.isdigit():
            from docx.shared import Pt
            run.font.size = Pt(int(size_str) / 2)
    if "bold" in params:
        run.bold = params["bold"]
    if "italic" in params:
        run.italic = params["italic"]
    if "color_rgb" in params:
        from docx.shared import RGBColor
        hex_str = params["color_rgb"].lstrip("#")
        if hex_str:
            run.font.color.rgb = RGBColor.from_string(hex_str)
        else:
            run.font.color.rgb = None
            from docx.oxml.ns import qn
            rpr = run._element.find(qn('w:rPr'))
            if rpr is not None:
                color_elem = rpr.find(qn('w:color'))
                if color_elem is not None:
                    rpr.remove(color_elem)
    if "color_theme" in params:
        from docx.oxml.ns import qn
        from docx.enum.dml import MSO_THEME_COLOR
        theme_map = {
            "DARK_1": MSO_THEME_COLOR.DARK_1,
            "LIGHT_1": MSO_THEME_COLOR.LIGHT_1,
            "DARK_2": MSO_THEME_COLOR.DARK_2,
            "LIGHT_2": MSO_THEME_COLOR.LIGHT_2,
            "ACCENT_1": MSO_THEME_COLOR.ACCENT_1,
            "ACCENT_2": MSO_THEME_COLOR.ACCENT_2,
            "ACCENT_3": MSO_THEME_COLOR.ACCENT_3,
            "ACCENT_4": MSO_THEME_COLOR.ACCENT_4,
            "ACCENT_5": MSO_THEME_COLOR.ACCENT_5,
            "ACCENT_6": MSO_THEME_COLOR.ACCENT_6,
            "HYPERLINK": MSO_THEME_COLOR.HYPERLINK,
            "FOLLOWED_HYPERLINK": MSO_THEME_COLOR.FOLLOWED_HYPERLINK,
            "TEXT_1": MSO_THEME_COLOR.DARK_1,
            "TEXT_2": MSO_THEME_COLOR.LIGHT_1,
            "BACKGROUND_1": MSO_THEME_COLOR.LIGHT_1,
            "BACKGROUND_2": MSO_THEME_COLOR.LIGHT_2,
        }
        theme_name = params["color_theme"].upper()
        if theme_name in theme_map:
            run.font.color.theme_color = theme_map[theme_name]

def _apply_set_font(paragraph, params):
    """Apply font settings to paragraph runs or specific matched text."""
    text_match = params.get("text_match")
    
    if not text_match:
        for run in paragraph.runs:
            _apply_font_attributes(run, params)
        return

    full_text = paragraph.text
    if text_match not in full_text:
        return
        
    match_index = params.get("match_index", "all")
    match_ranges = []
    start = 0
    while True:
        idx = full_text.find(text_match, start)
        if idx == -1:
            break
        match_ranges.append((idx, idx + len(text_match)))
        start = idx + len(text_match)
        
    if str(match_index) != "all":
        try:
            mi = int(match_index)
            if mi < 0 or mi >= len(match_ranges):
                return
            match_ranges = [match_ranges[mi]]
        except ValueError:
            pass

    if not match_ranges:
        return

    original_runs_info = []
    current_idx = 0
    for r in paragraph.runs:
        r_len = len(r.text)
        original_runs_info.append({
            'text': r.text,
            'start': current_idx,
            'end': current_idx + r_len,
            'run': r
        })
        current_idx += r_len

    paragraph.clear()

    for r_info in original_runs_info:
        r_start = r_info['start']
        r_end = r_info['end']
        r_text = r_info['text']
        
        intersections = []
        for m_start, m_end in match_ranges:
            if m_start < r_end and m_end > r_start:
                intersections.append((max(r_start, m_start), min(r_end, m_end)))
                
        if not intersections:
            new_run = paragraph.add_run(r_text)
            _copy_run_format(r_info['run'], new_run)
            continue
            
        curr_offset = 0
        for i_start, i_end in intersections:
            rel_start = i_start - r_start
            rel_end = i_end - r_start
            
            if rel_start > curr_offset:
                new_run = paragraph.add_run(r_text[curr_offset:rel_start])
                _copy_run_format(r_info['run'], new_run)
                
            new_run = paragraph.add_run(r_text[rel_start:rel_end])
            _copy_run_format(r_info['run'], new_run)
            _apply_font_attributes(new_run, params)
            
            curr_offset = rel_end
            
        if curr_offset < len(r_text):
            new_run = paragraph.add_run(r_text[curr_offset:])
            _copy_run_format(r_info['run'], new_run)

def _apply_set_paragraph_format(paragraph, params):
    """Apply paragraph format settings."""
    pf = paragraph.paragraph_format
    if "alignment" in params:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        pf.alignment = align_map.get(params["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
    if "line_spacing" in params:
        pf.line_spacing = params["line_spacing"]
    if "first_line_indent" in params:
        indent_str = params["first_line_indent"]
        if indent_str.endswith("em"):
            from docx.shared import Emu
            pf.first_line_indent = Emu(int(float(indent_str.rstrip("em")) * 914400 / 2))
        elif indent_str.endswith("pt"):
            from docx.shared import Pt
            pf.first_line_indent = Pt(float(indent_str.rstrip("pt")))
    if "space_before" in params:
        space_str = params["space_before"]
        if space_str.endswith("pt"):
            from docx.shared import Pt
            pf.space_before = Pt(float(space_str.rstrip("pt")))
    if "space_after" in params:
        space_str = params["space_after"]
        if space_str.endswith("pt"):
            from docx.shared import Pt
            pf.space_after = Pt(float(space_str.rstrip("pt")))

def _apply_insert_paragraph(doc, params):
    """Insert a new paragraph before or after the anchor paragraph."""
    content = params.get('content')
    if content is None:
        print("[WARNING] op=insert_paragraph: missing content, skipping.")
        return

    before_id = params.get('before')
    after_id = params.get('after')

    if bool(before_id) == bool(after_id):
        print("[WARNING] op=insert_paragraph: must provide exactly one of 'before' or 'after', skipping.")
        return

    target_id = before_id or after_id
    try:
        idx = int(target_id.replace('p', ''))
        if idx < 0 or idx >= len(doc.paragraphs):
            print(f"[WARNING] op=insert_paragraph: paragraph index {idx} out of bounds, skipping.")
            return
    except (ValueError, AttributeError):
        print("[WARNING] op=insert_paragraph: invalid target format, skipping.")
        return

    anchor_p = doc.paragraphs[idx]
    style_name = params.get('style')

    if before_id:
        new_p = anchor_p.insert_paragraph_before(content)
        if style_name:
            new_p.style = style_name
        elif idx == 0:
            new_p.style = 'Normal'
        else:
            new_p.style = anchor_p.style.name
    else:
        new_p = doc.add_paragraph(content)
        anchor_p._p.addnext(new_p._p)
        new_p.style = style_name if style_name else anchor_p.style.name


def _apply_delete_paragraph(doc, params):
    """Delete the paragraph at the given index."""
    target = params.get('target')
    if not target:
        print("[WARNING] op=delete_paragraph: missing target, skipping.")
        return
    try:
        idx = int(target.replace('p', ''))
        if idx < 0 or idx >= len(doc.paragraphs):
            print(f"[WARNING] op=delete_paragraph: paragraph index {idx} out of bounds, skipping.")
            return
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
    except (ValueError, AttributeError):
        print("[WARNING] op=delete_paragraph: invalid target format, skipping.")


def _apply_modify_cell(doc, params):
    t_idx, r_idx, c_idx = params.get('table'), params.get('row'), params.get('col')
    if any(x is None for x in [t_idx, r_idx, c_idx, params.get('content')]):
        print("[WARNING] op=modify_cell: missing required parameters, skipping.")
        return
    try:
        table = doc.tables[t_idx]
        table.cell(r_idx, c_idx).text = params['content']
    except IndexError:
        print("[WARNING] op=modify_cell: index out of bounds, skipping.")


def _apply_insert_row(doc, params):
    t_idx = params.get('table')
    before, after = params.get('before'), params.get('after')
    if t_idx is None or (before is None and after is None) or (before is not None and after is not None):
        print("[WARNING] op=insert_row: invalid parameters, skipping.")
        return
    try:
        table = doc.tables[t_idx]
        target_row_idx = before if before is not None else after
        target_row = table.rows[target_row_idx]
        new_row = table.add_row()
        if before is not None:
            target_row._tr.addprevious(new_row._tr)
        else:
            target_row._tr.addnext(new_row._tr)
    except IndexError:
        print("[WARNING] op=insert_row: index out of bounds, skipping.")


def _apply_insert_table(doc, params):
    before_id, after_id = params.get('before'), params.get('after')
    rows, cols = params.get('rows'), params.get('cols')
    if (before_id is None and after_id is None) or (before_id and after_id) or rows is None or cols is None:
        print("[WARNING] op=insert_table: invalid parameters, skipping.")
        return
    
    target_id = before_id or after_id
    try:
        idx = int(target_id.replace('p', ''))
        if idx < 0 or idx >= len(doc.paragraphs):
            print(f"[WARNING] op=insert_table: paragraph index {idx} out of bounds, skipping.")
            return
        anchor_p = doc.paragraphs[idx]
        new_table = doc.add_table(rows=rows, cols=cols)
        if before_id:
            anchor_p._p.addprevious(new_table._tbl)
        else:
            anchor_p._p.addnext(new_table._tbl)
    except (ValueError, AttributeError, IndexError):
        print("[WARNING] op=insert_table: invalid target format or out of bounds, skipping.")


def _apply_set_header_footer(doc, params):
    sec_idx = params.get('section', 0)
    try:
        section = doc.sections[sec_idx]
        even = params.get('even_page', False)
        header_text = params.get('header')
        footer_text = params.get('footer')
        
        target_header = section.even_page_header if even else section.header
        target_footer = section.even_page_footer if even else section.footer
        
        if header_text is not None:
            if not target_header.paragraphs:
                target_header.add_paragraph()
            target_header.paragraphs[0].text = header_text
            
        if footer_text is not None:
            if not target_footer.paragraphs:
                target_footer.add_paragraph()
            target_footer.paragraphs[0].text = footer_text
    except IndexError:
        print("[WARNING] op=set_header_footer: section index out of bounds, skipping.")


def _apply_insert_page_break(doc, params):
    before_id, after_id = params.get('before'), params.get('after')
    if (not before_id and not after_id) or (before_id and after_id):
        print("[WARNING] op=insert_page_break: invalid parameters, skipping.")
        return
        
    target_id = before_id or after_id
    try:
        idx = int(target_id.replace('p', ''))
        anchor_p = doc.paragraphs[idx]
        new_p = doc.add_paragraph()
        new_p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        if before_id:
            anchor_p._p.addprevious(new_p._p)
        else:
            anchor_p._p.addnext(new_p._p)
    except (ValueError, AttributeError, IndexError):
        print("[WARNING] op=insert_page_break: invalid target format, skipping.")


def _parse_unit(val_str):
    from docx.shared import Cm, Pt
    if val_str.endswith('cm'):
        return Cm(float(val_str.rstrip('cm')))
    elif val_str.endswith('pt'):
        return Pt(float(val_str.rstrip('pt')))
    return None


def _apply_apply_style(doc, params):
    target = params.get('target')
    style_name = params.get('style')
    if not target or not style_name:
        print("[WARNING] op=apply_style: missing parameters, skipping.")
        return
    try:
        idx = int(target.replace('p', ''))
        if idx < 0 or idx >= len(doc.paragraphs):
            print(f"[WARNING] op=apply_style: paragraph index {idx} out of bounds, skipping.")
            return
        if style_name in doc.styles:
            para = doc.paragraphs[idx]
            para.style = style_name
            if params.get('clear_run_formats'):
                from docx.oxml.ns import qn
                for run in para.runs:
                    run.font.size = None
                    run.bold = None
                    run.italic = None
                    run.font.name = None
                    # Clear both rgb and theme color overrides
                    run.font.color.rgb = None
                    run.font.color.theme_color = None
                    # Clear XML-level overrides (East Asia font, color element)
                    rpr = run._element.find(qn('w:rPr'))
                    if rpr is not None:
                        rfonts = rpr.find(qn('w:rFonts'))
                        if rfonts is not None:
                            rfonts.attrib.pop(qn('w:eastAsia'), None)
                        # Remove w:color element to fully clear color override
                        color_elem = rpr.find(qn('w:color'))
                        if color_elem is not None:
                            rpr.remove(color_elem)
        else:
            print(f"[WARNING] op=apply_style: style '{style_name}' not found, skipping.")
    except (ValueError, AttributeError, IndexError):
        print("[WARNING] op=apply_style: invalid target format or index, skipping.")


def _apply_set_page_setup(doc, params):
    sec_idx = params.get('section', 0)
    try:
        section = doc.sections[sec_idx]
        for margin in ['margin_top', 'margin_bottom', 'margin_left', 'margin_right']:
            if margin in params:
                val = _parse_unit(params[margin])
                if val:
                    if margin == 'margin_top': section.top_margin = val
                    elif margin == 'margin_bottom': section.bottom_margin = val
                    elif margin == 'margin_left': section.left_margin = val
                    elif margin == 'margin_right': section.right_margin = val
        
        orientation = params.get('orientation')
        if orientation == 'landscape':
            section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        elif orientation == 'portrait':
            section.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
    except IndexError:
        print("[WARNING] op=set_page_setup: section index out of bounds, skipping.")


def _apply_update_style_definition(doc, params):
    """Overwrite a built-in style's font/paragraph format with fingerprint values."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    style_name = params.get('style')
    fp = params.get('fingerprint', {})
    if not style_name or not fp:
        print("[WARNING] op=update_style_definition: missing style or fingerprint, skipping.")
        return
    if style_name not in doc.styles:
        print(f"[WARNING] op=update_style_definition: style '{style_name}' not found, skipping.")
        return

    style = doc.styles[style_name]

    # Font size
    size_str = fp.get('size')
    if size_str:
        try:
            style.font.size = Pt(float(size_str.rstrip('pt')))
        except (ValueError, AttributeError):
            pass

    # Bold / italic
    if fp.get('bold') is not None:
        style.font.bold = fp['bold']
    if fp.get('italic') is not None:
        style.font.italic = fp['italic']

    # Western font (name)
    font_name = fp.get('font')
    if font_name:
        style.font.name = font_name
        # Also set East Asia (Chinese) font via XML
        from docx.oxml.ns import qn
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rfonts)
        # Clear theme font attributes to avoid precedence issues
        for attr in list(rfonts.attrib):
            localname = attr.split('}')[-1] if '}' in attr else attr
            if localname in ('eastAsiaTheme', 'asciiTheme', 'hAnsiTheme', 'cstheme', 'majorHAnsi', 'majorEastAsia', 'majorBidi', 'minorBidi'):
                del rfonts.attrib[attr]
        rfonts.set(qn('w:eastAsia'), font_name)
        rfonts.set(qn('w:ascii'), font_name)
        rfonts.set(qn('w:hAnsi'), font_name)

    # Color
    # python-docx's ColorFormat clears the entire w:color element when setting
    # rgb=None or theme_color=None, so we must manipulate XML directly to set
    # one attribute while clearing the other.
    color = fp.get('color')
    if color:
        from docx.oxml.ns import qn
        rpr = style._element.get_or_add_rPr()
        # Remove existing w:color element to start clean
        for old_color in rpr.findall(qn('w:color')):
            rpr.remove(old_color)
        if color.startswith('rgb:'):
            from docx.shared import RGBColor
            hex_str = color[4:].lstrip('#')
            if hex_str:
                color_elem = rpr.makeelement(qn('w:color'), {
                    qn('w:val'): hex_str,
                })
                rpr.append(color_elem)
        elif color.startswith('theme:'):
            from docx.enum.dml import MSO_THEME_COLOR
            theme_map = {
                "DARK_1": MSO_THEME_COLOR.DARK_1,
                "LIGHT_1": MSO_THEME_COLOR.LIGHT_1,
                "DARK_2": MSO_THEME_COLOR.DARK_2,
                "LIGHT_2": MSO_THEME_COLOR.LIGHT_2,
                "ACCENT_1": MSO_THEME_COLOR.ACCENT_1,
                "ACCENT_2": MSO_THEME_COLOR.ACCENT_2,
                "ACCENT_3": MSO_THEME_COLOR.ACCENT_3,
                "ACCENT_4": MSO_THEME_COLOR.ACCENT_4,
                "ACCENT_5": MSO_THEME_COLOR.ACCENT_5,
                "ACCENT_6": MSO_THEME_COLOR.ACCENT_6,
                "HYPERLINK": MSO_THEME_COLOR.HYPERLINK,
                "FOLLOWED_HYPERLINK": MSO_THEME_COLOR.FOLLOWED_HYPERLINK,
                "TEXT_1": MSO_THEME_COLOR.DARK_1,
                "TEXT_2": MSO_THEME_COLOR.LIGHT_1,
                "BACKGROUND_1": MSO_THEME_COLOR.LIGHT_1,
                "BACKGROUND_2": MSO_THEME_COLOR.LIGHT_2,
            }
            theme_name = color[6:].upper()
            if theme_name in theme_map:
                theme_val = theme_map[theme_name]
                color_elem = rpr.makeelement(qn('w:color'), {
                    qn('w:themeColor'): theme_val.xml_value,
                })
                rpr.append(color_elem)
    elif color is None or color == '':
        # Clear color when fingerprint has no color (use theme default)
        from docx.oxml.ns import qn
        rpr = style._element.find(qn('w:rPr'))
        if rpr is not None:
            for old_color in rpr.findall(qn('w:color')):
                rpr.remove(old_color)

    # Alignment
    align_str = fp.get('align')
    if align_str:
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if align_str in align_map:
            style.paragraph_format.alignment = align_map[align_str]

    # Spacing
    sb = fp.get('space_before')
    if sb:
        try:
            style.paragraph_format.space_before = Pt(float(sb.rstrip('pt')))
        except (ValueError, AttributeError):
            pass

    sa = fp.get('space_after')
    if sa:
        try:
            style.paragraph_format.space_after = Pt(float(sa.rstrip('pt')))
        except (ValueError, AttributeError):
            pass

    # Line spacing (multiplier only)
    ls = fp.get('line_spacing')
    if ls is not None:
        style.paragraph_format.line_spacing = ls

    # First line indent
    fli = fp.get('first_line_indent')
    if fli:
        try:
            style.paragraph_format.first_line_indent = Pt(float(fli.rstrip('pt')))
        except (ValueError, AttributeError):
            pass

    # Visibility: make sure heading styles are visible in the style gallery
    # (some source documents have Heading 2/3 set to hidden=True)
    if style_name.startswith('Heading '):
        style.hidden = False


def _insert_tbl_pr_element(tblPr, tag_name):
    """
    Safely get or insert a child element into w:tblPr maintaining strict OOXML order.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # 查找已有节点
    existing = tblPr.find(qn(f'w:{tag_name}'))
    if existing is not None:
        # 清空已有节点的子节点，保留节点本身
        for child in list(existing):
            existing.remove(child)
        return existing
        
    # 定义 w:tblPr 子元素严格 OOXML 顺序
    TBL_PR_ORDER = [
        'tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
        'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd',
        'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook', 'activeRecord'
    ]
    
    if tag_name not in TBL_PR_ORDER:
        # 保守回退：直接 append
        new_elem = OxmlElement(f'w:{tag_name}')
        tblPr.append(new_elem)
        return new_elem
        
    target_idx = TBL_PR_ORDER.index(tag_name)
    
    # 找到应该插入的位置
    insert_before_elem = None
    for child in tblPr:
        # 获取子节点的 tag（不含 namespace）
        child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if child_tag in TBL_PR_ORDER:
            child_idx = TBL_PR_ORDER.index(child_tag)
            if child_idx > target_idx:
                insert_before_elem = child
                break
                
    new_elem = OxmlElement(f'w:{tag_name}')
    if insert_before_elem is not None:
        insert_before_elem.addprevious(new_elem)
    else:
        tblPr.append(new_elem)
        
    return new_elem


def _insert_tc_pr_element(tcPr, tag_name):
    """
    Safely get or insert a child element into w:tcPr maintaining strict OOXML order.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # 查找已有节点
    existing = tcPr.find(qn(f'w:{tag_name}'))
    if existing is not None:
        # 清空已有节点的属性和子节点，保留节点本身
        existing.attrib.clear()
        for child in list(existing):
            existing.remove(child)
        return existing
        
    # 定义 w:tcPr 子元素严格 OOXML 顺序
    TC_PR_ORDER = [
        'cnfStyle', 'tcW', 'gridSpan', 'hMerge', 'vMerge', 'tcBorders', 'shd',
        'noWrap', 'tcMar', 'textDirection', 'fitText', 'vAlign', 'wmlFlow', 'tcFitText'
    ]
    
    if tag_name not in TC_PR_ORDER:
        new_elem = OxmlElement(f'w:{tag_name}')
        tcPr.append(new_elem)
        return new_elem
        
    target_idx = TC_PR_ORDER.index(tag_name)
    
    insert_before_elem = None
    for child in tcPr:
        child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if child_tag in TC_PR_ORDER:
            child_idx = TC_PR_ORDER.index(child_tag)
            if child_idx > target_idx:
                insert_before_elem = child
                break
                
    new_elem = OxmlElement(f'w:{tag_name}')
    if insert_before_elem is not None:
        insert_before_elem.addprevious(new_elem)
    else:
        tcPr.append(new_elem)
        
    return new_elem


def _apply_table_style(doc, params):
    """
    Apply table style parameters (borders, shading, margins, cell texts) to a table.
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    target = params.get("target")
    table_role = params.get("table_role", {})
    if not target or not table_role:
        print("[WARNING] op=apply_table_style: missing target or table_role, skipping.")
        return

    try:
        t_idx = int(target.replace("t", ""))
        if t_idx < 0 or t_idx >= len(doc.tables):
            print(f"[WARNING] op=apply_table_style: table index {t_idx} out of bounds, skipping.")
            return
    except (ValueError, AttributeError):
        print("[WARNING] op=apply_table_style: invalid target format, skipping.")
        return

    table = doc.tables[t_idx]

    # 1. 应用 tbl_style (Word 内置样式名)
    tbl_style_name = table_role.get("tbl_style")
    if tbl_style_name:
        if tbl_style_name in doc.styles:
            style = doc.styles[tbl_style_name]
            if style.type == WD_STYLE_TYPE.TABLE:
                table.style = style
            else:
                print(f"[WARNING] op=apply_table_style: style '{tbl_style_name}' is not a TABLE style, skipping style mapping.")
        else:
            print(f"[WARNING] op=apply_table_style: style '{tbl_style_name}' not found, skipping style mapping.")

    # 2. 写入六方向边框 (tblBorders)
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    border_info = table_role.get("border")
    if border_info:
        tblBorders = _insert_tbl_pr_element(tblPr, "tblBorders")

        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b_data = border_info.get(side)
            if b_data and isinstance(b_data, dict):
                val = b_data.get("val")
                sz = b_data.get("sz")
                color = b_data.get("color")
                if val:
                    elem = OxmlElement(f"w:{side}")
                    elem.set(qn("w:val"), val)
                    if sz is not None:
                        elem.set(qn("w:sz"), str(sz))
                    if color:
                        elem.set(qn("w:color"), color)
                    tblBorders.append(elem)

    # 3. 写入单元格内边距 (tblCellMar)
    cell_margin = table_role.get("cell_margin")
    if cell_margin:
        tblCellMar = _insert_tbl_pr_element(tblPr, "tblCellMar")

        for side in ("top", "bottom", "left", "right"):
            val_pt_str = cell_margin.get(side)
            if val_pt_str:
                try:
                    val_pt = float(val_pt_str.rstrip("pt"))
                    val_dxa = int(round(val_pt * 20.0))
                    elem = OxmlElement(f"w:{side}")
                    elem.set(qn("w:w"), str(val_dxa))
                    elem.set(qn("w:type"), "dxa")
                    tblCellMar.append(elem)
                except ValueError:
                    pass

    # 4. 逐行应用底纹和文字格式 (包含合并单元格排重逻辑)
    has_header = table_role.get("structure", {}).get("has_header_row", False)
    header_shading = table_role.get("shading", {}).get("header")
    body_shading = table_role.get("shading", {}).get("body")
    header_text_format = table_role.get("header_text")
    body_text_format = table_role.get("body_text")

    processed_cells = set()

    for r_idx, row in enumerate(table.rows):
        is_header_row = (r_idx == 0 and has_header)
        shading_color = header_shading if is_header_row else body_shading
        text_format = header_text_format if is_header_row else body_text_format

        for cell in row.cells:
            # 合并单元格排重：使用底层 _tc 元素识别
            cell_key = cell._tc
            if cell_key in processed_cells:
                continue
            processed_cells.add(cell_key)

            # 4.1 底纹应用
            tcPr = cell._tc.get_or_add_tcPr()
            # 原地获取或创建 w:shd 节点写入，防止 XML 损坏并保证子节点严格顺序
            shd = _insert_tc_pr_element(tcPr, "shd")
            
            if shading_color:
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), shading_color)
            else:
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "auto")

            # 4.2 文字格式应用
            if text_format:
                # 转换格式参数为 _apply_set_font 和 _apply_set_paragraph_format 所需结构
                font_params = {}
                para_params = {}

                if text_format.get("font"):
                    font_params["name"] = text_format["font"]
                    font_params["east_asia"] = text_format["font"]
                if text_format.get("size"):
                    font_params["size"] = text_format["size"]
                if text_format.get("bold") is not None:
                    font_params["bold"] = text_format["bold"]

                if text_format.get("align"):
                    para_params["alignment"] = text_format["align"]

                for paragraph in cell.paragraphs:
                    # 如果段落没有 runs，添加一个空 run 以便应用字体格式
                    if not paragraph.runs and paragraph.text:
                        paragraph.add_run(paragraph.text)
                    elif not paragraph.runs:
                        paragraph.add_run()

                    # 应用段落级别格式
                    if para_params:
                        _apply_set_paragraph_format(paragraph, para_params)
                    # 应用 run 级别字体格式
                    if font_params:
                        _apply_set_font(paragraph, font_params)


def _parse_cell_target(target: str):
    """
    解析 'tNrNcN' 格式的单元格 target，返回 (table_idx, row_idx, col_idx)。
    格式非法时抛出 ValueError；索引越界时由调用方抛出 IndexError。
    """
    m = re.fullmatch(r't(\d+)r(\d+)c(\d+)', target)
    if not m:
        raise ValueError(
            f"Invalid cell target: {target!r}. Expected format: tNrNcN (e.g. t0r1c2)"
        )
    return int(m.group(1)), int(m.group(2)), int(m.group(3))


def _apply_replace_cell_text(doc, op):
    """
    replace_cell_text: 在指定单元格内执行文本替换。
    op 必须包含: target (tNrNcN), find (str), replace (str)
    """
    t_idx, r_idx, c_idx = _parse_cell_target(op['target'])
    tables = doc.tables
    if t_idx >= len(tables):
        raise IndexError(
            f"Table index {t_idx} out of range (document has {len(tables)} table(s))"
        )
    rows = tables[t_idx].rows
    if r_idx >= len(rows):
        raise IndexError(
            f"Row index {r_idx} out of range (table t{t_idx} has {len(rows)} row(s))"
        )
    cells = rows[r_idx].cells
    if c_idx >= len(cells):
        raise IndexError(
            f"Col index {c_idx} out of range (row t{t_idx}r{r_idx} has {len(cells)} cell(s))"
        )
    cell = cells[c_idx]
    find_txt = op['find']
    repl_txt = op['replace']
    for para in cell.paragraphs:
        _replace_text_in_runs(para, find_txt, repl_txt)


def _apply_rewrite_cell(doc, op):
    """
    rewrite_cell: 整体重写指定单元格的纯文本内容。
    op 必须包含: target (tNrNcN), content (str)
    """
    t_idx, r_idx, c_idx = _parse_cell_target(op['target'])
    tables = doc.tables
    if t_idx >= len(tables):
        raise IndexError(
            f"Table index {t_idx} out of range (document has {len(tables)} table(s))"
        )
    rows = tables[t_idx].rows
    if r_idx >= len(rows):
        raise IndexError(
            f"Row index {r_idx} out of range (table t{t_idx} has {len(rows)} row(s))"
        )
    cells = rows[r_idx].cells
    if c_idx >= len(cells):
        raise IndexError(
            f"Col index {c_idx} out of range (row t{t_idx}r{r_idx} has {len(cells)} cell(s))"
        )
    cell = cells[c_idx]
    # 清空所有段落后写入新内容
    for para in cell.paragraphs:
        para.text = ''
    if cell.paragraphs:
        cell.paragraphs[0].text = op['content']
    else:
        cell.add_paragraph(op['content'])


def _backup_file(filepath):
    """Create a .bak backup of the original file before modifying."""
    bak_path = filepath + '.bak'
    shutil.copy2(filepath, bak_path)
    return bak_path

def apply_operations(filepath, ops, outpath=None):
    doc = docx.Document(filepath)
    if outpath is None:
        outpath = filepath

    # Backup original before any modification
    if os.path.exists(filepath):
        _backup_file(filepath)

    for op in ops:
        if op['op'] == 'rewrite_paragraph':
            idx = int(op['target'].replace('p', ''))
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                saved_style = para.style
                para.text = op['content']
                para.style = saved_style
        elif op['op'] == 'replace_text':
            target = op.get('target', 'all')
            find_txt = op['find']
            repl_txt = op['replace']
            if target == 'all':
                for p in doc.paragraphs:
                    _replace_text_in_runs(p, find_txt, repl_txt)
            else:
                idx = int(target.replace('p', ''))
                if idx < len(doc.paragraphs):
                    _replace_text_in_runs(doc.paragraphs[idx], find_txt, repl_txt)
        elif op['op'] == 'set_font':
            idx = int(op['target'].replace('p', ''))
            if idx < len(doc.paragraphs):
                _apply_set_font(doc.paragraphs[idx], op)
        elif op['op'] == 'set_paragraph_format':
            idx = int(op['target'].replace('p', ''))
            if idx < len(doc.paragraphs):
                _apply_set_paragraph_format(doc.paragraphs[idx], op)
        elif op['op'] == 'insert_paragraph':
            _apply_insert_paragraph(doc, op)
        elif op['op'] == 'delete_paragraph':
            _apply_delete_paragraph(doc, op)
        elif op['op'] == 'modify_cell':
            _apply_modify_cell(doc, op)
        elif op['op'] == 'insert_row':
            _apply_insert_row(doc, op)
        elif op['op'] == 'insert_table':
            _apply_insert_table(doc, op)
        elif op['op'] == 'set_header_footer':
            _apply_set_header_footer(doc, op)
        elif op['op'] == 'insert_page_break':
            _apply_insert_page_break(doc, op)
        elif op['op'] == 'apply_style':
            _apply_apply_style(doc, op)
        elif op['op'] == 'apply_table_style':
            _apply_table_style(doc, op)
        elif op['op'] == 'set_page_setup':
            _apply_set_page_setup(doc, op)
        elif op['op'] == 'update_style_definition':
            _apply_update_style_definition(doc, op)
        elif op['op'] == 'replace_cell_text':
            _apply_replace_cell_text(doc, op)
        elif op['op'] == 'rewrite_cell':
            _apply_rewrite_cell(doc, op)

    doc.save(outpath)

if __name__ == "__main__":
    filepath = sys.argv[1]
    with open(sys.argv[2], 'r', encoding='utf-8') as f:
        ops = json.load(f)
    outpath = sys.argv[3] if len(sys.argv) > 3 else filepath
    apply_operations(filepath, ops, outpath)
