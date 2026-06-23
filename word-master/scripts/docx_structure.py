"""
DOCX 标题编号管理：添加/移除编号前缀。

与 markdown-master/scripts/structure.py 的 numbering add/remove 对称，
但针对 docx 按 style.name 识别标题（而非 md 的 # 前缀），
编号写为文本前缀（非 Word 原生 numPr）。
"""
import re
import shutil
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from _numbering_render import (
    parse_template, format_template, load_config, save_config,
)

import docx

_HEADING_RE = re.compile(r"^Heading (\d+)$")
_HEADING_CN_RE = re.compile(r"^标题\s*(\d+)$")


def _detect_heading_level(paragraph):
    """按 paragraph.style.name 判定标题级别（1-9），非标题返回 None。"""
    name = paragraph.style.name if paragraph.style else ""
    for pat in (_HEADING_RE, _HEADING_CN_RE):
        m = pat.match(name)
        if m:
            level = int(m.group(1))
            if 1 <= level <= 9:
                return level
    return None


_NUMBER_PREFIX_RE = re.compile(
    r"^("
    r"第[\d]+章\s*"
    r"|[\d]+[、，]\s*"
    r"|（[\d]+）\s*"
    r"|[IVXLCDM]+\s+"
    r"|[ivxlcdm]+\s+"
    r"|[A-Z]\s+"
    r"|[a-z]\s+"
    r"|[\d]+(?:\.[\d]+)*\.?\s+"
    r"|[零一二三四五六七八九十百千]+章\s*"
    r")"
)


def _strip_prefix(text):
    """剥除标题文本开头的编号前缀（只剥一次）。"""
    return _NUMBER_PREFIX_RE.sub("", text, count=1)


def _save_docx(doc, docx_path, output_path):
    """保存 docx。output_path 给定则写新文件；否则覆盖原文件并生成 .bak。"""
    if output_path:
        doc.save(output_path)
    else:
        shutil.copy2(docx_path, docx_path + ".bak")
        doc.save(docx_path)


def cmd_numbering_add(docx_path, level_templates, start_from=1,
                      output_path=None, save_config_path=None):
    """
    给 docx 的标题段落添加编号前缀。

    level_templates: dict[1..6] = str | None，空串/None 表示该级不编号。
    start_from: h1 起始编号。
    output_path: 给定则输出到新文件；None 则覆盖原文件并生成 .bak。
    save_config_path: 给定则把当前模板存成 YAML（同时仍执行编号）。
    """
    doc = docx.Document(docx_path)
    counters = [0] * 6
    if level_templates.get(1):
        counters[0] = start_from - 1

    for para in doc.paragraphs:
        level = _detect_heading_level(para)
        if level is None or level > 6:
            continue
        template = level_templates.get(level)
        if not template:
            continue
        counters[level - 1] += 1
        for j in range(level, 6):
            counters[j] = 0

        tokens = parse_template(template)
        prefix = format_template(tokens, counters)

        if not para.runs:
            para.add_run(prefix)
        else:
            stripped = _strip_prefix(para.runs[0].text)
            para.runs[0].text = prefix + stripped

    _save_docx(doc, docx_path, output_path)

    if save_config_path:
        cfg = dict(level_templates)
        cfg["start_from"] = start_from
        save_config(save_config_path, cfg)