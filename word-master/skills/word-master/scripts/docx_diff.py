import docx
import re
import sys
import difflib
from dataclasses import dataclass, field
from typing import List, Optional

@dataclass
class ParagraphBlock:
    index: int
    text: str
    style: str
    overrides: List[str] = field(default_factory=list)
    raw: str = ""  # Original markdown line

def parse_markdown_blocks(md: str) -> List[ParagraphBlock]:
    """
    Parse rich markdown into structured ParagraphBlock list.
    Extracts: paragraph index [pN], text, style {StyleName}, overrides.
    """
    blocks = []
    # Match: ### [p0] Text {StyleName} or ### [p0] Text {StyleName | override: ...}
    pattern = r'### \[p(\d+)\] (.+?) \{(.+?)\}'
    for line in md.split('\n'):
        match = re.search(pattern, line)
        if match:
            idx = int(match.group(1))
            text = match.group(2).strip()
            style_part = match.group(3)
            # Split style and overrides
            if '|' in style_part:
                style, override_part = style_part.split('|', 1)
                overrides = [o.strip() for o in override_part.replace('override:', '').split(';')]
                overrides = [o for o in overrides if o]
            else:
                style = style_part
                overrides = []
            blocks.append(ParagraphBlock(
                index=idx,
                text=text,
                style=style.strip(),
                overrides=overrides,
                raw=line
            ))
    return blocks

def generate_diff(path1, path2):
    d1 = docx.Document(path1)
    d2 = docx.Document(path2)
    
    texts1 = [p.text for p in d1.paragraphs if p.text.strip()]
    texts2 = [p.text for p in d2.paragraphs if p.text.strip()]
    
    diff = difflib.unified_diff(texts1, texts2, lineterm='', n=0)
    return '\n'.join(diff)

def align_paragraphs(blocks1: List[ParagraphBlock], blocks2: List[ParagraphBlock]) -> List[tuple]:
    """
    Align paragraphs between two documents.
    Returns: List[(block1_or_None, block2_or_None, status)]
    status: 'same', 'modified', 'new', 'deleted'
    """
    aligned = []

    # Build a map from original index to block for blocks2
    idx2_map = {b.index: b for b in blocks2}

    for i in range(len(blocks1)):
        b1 = blocks1[i]
        b2 = idx2_map.get(b1.index)

        if b2 is None:
            # Original index not in blocks2 -> deleted
            aligned.append((b1, None, 'deleted'))
        elif b1.text == b2.text and b1.style == b2.style and b1.overrides == b2.overrides:
            # Same
            aligned.append((b1, b2, 'same'))
        else:
            # Modified
            aligned.append((b1, b2, 'modified'))

    # Remaining paragraphs in blocks2 (at indices beyond len(blocks1)) are new
    for b2 in blocks2:
        if b2.index >= len(blocks1):
            aligned.append((None, b2, 'new'))

    return aligned

@dataclass
class DiffResult:
    type: str  # 'same', 'content', 'format', 'new', 'deleted'
    original: Optional[ParagraphBlock]
    modified: Optional[ParagraphBlock]
    fields: List[dict] = field(default_factory=list)

def compute_field_diff(block1: Optional[ParagraphBlock], block2: Optional[ParagraphBlock]) -> DiffResult:
    """Compute field-level differences between two paragraph blocks."""
    if block1 is None and block2 is not None:
        return DiffResult(type='new', original=None, modified=block2, fields=[])
    if block1 is not None and block2 is None:
        return DiffResult(type='deleted', original=block1, modified=None, fields=[])
    if block1 is None and block2 is None:
        return DiffResult(type='same', original=None, modified=None, fields=[])

    fields = []
    change_type = 'same'

    # Style change
    if block1.style != block2.style:
        fields.append({'name': 'style', 'orig': block1.style, 'modif': block2.style})
        change_type = 'format'

    # Overrides change
    if block1.overrides != block2.overrides:
        fields.append({'name': 'overrides', 'orig': block1.overrides, 'modif': block2.overrides})
        change_type = 'format'

    # Content change
    if block1.text != block2.text:
        fields.append({'name': 'content', 'orig': block1.text, 'modif': block2.text})
        change_type = 'content'

    return DiffResult(type=change_type, original=block1, modified=block2, fields=fields)

if __name__ == "__main__":
    print(generate_diff(sys.argv[1], sys.argv[2]))
