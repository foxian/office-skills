import os
import sys
import docx
from collections import Counter
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def get_table_cols(table):
    """
    提取物理列数。
    尝试读取 w:tblGrid/w:gridCol，若读取不到，则使用 max(len(row.cells) for row in table.rows) 容错。
    """
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn('w:gridCol'))
        if gridCols:
            return len(gridCols)
    # 容错：使用最大的行单元格数
    if not table.rows:
        return 0
    return max(len(row.cells) for row in table.rows)

def get_effective_shading(cell):
    """
    判定有效底纹。获取单元格底纹配置，去除 w:val="clear" 时填充为常见默认白色/透明等（如 auto, FFFFFF）的底纹。
    """
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    
    val = shd.get(qn('w:val'))
    fill = shd.get(qn('w:fill'))
    
    if not fill:
        return None
        
    fill_upper = fill.upper()
    # 常见白色、自动、透明底纹均视作无有效底纹
    if fill_upper in ('AUTO', 'FFFFFF', 'FFF', 'CLEAR', 'NONE', 'NIL', '00000000'):
        return None
        
    if val in ('nil', 'none'):
        return None
        
    return fill_upper

def detect_has_header_row(table):
    """
    智能判定是否包含表头（三级检测）。
    三级校验：存在 w:tblHeader 标签 ➡️ 首行全部具有有效底纹 ➡️ 首行所有文本均加粗
    """
    if not table.rows:
        return False
        
    first_row = table.rows[0]
    
    # 1. 校验是否存在 w:tblHeader 标签
    trPr = first_row._tr.find(qn('w:trPr'))
    if trPr is not None:
        tblHeader = trPr.find(qn('w:tblHeader'))
        if tblHeader is not None:
            return True
            
    # 2. 校验首行全部具有有效底纹
    cells = first_row.cells
    if cells:
        all_have_shading = True
        for cell in cells:
            if get_effective_shading(cell) is None:
                all_have_shading = False
                break
        if all_have_shading:
            return True
            
    # 3. 校验首行所有文本均加粗
    def cell_is_bold(c):
        has_any_text = False
        all_bold = True
        for p in c.paragraphs:
            if p.text.strip():
                has_any_text = True
                p_has_text = False
                p_all_bold = True
                for r in p.runs:
                    if r.text.strip():
                        p_has_text = True
                        if r.bold is not True:
                            p_all_bold = False
                if p_has_text:
                    if not p_all_bold:
                        return False
                else:
                    # 没有 runs，但段落文本不为空，退化为样式继承校验
                    style_bold = False
                    style = p.style
                    for _ in range(10):
                        if style is None:
                            break
                        style_font = style.font if hasattr(style, 'font') else None
                        if style_font and style_font.bold is not None:
                            style_bold = style_font.bold
                            break
                        style = style.base_style if hasattr(style, 'base_style') else None
                    if not style_bold:
                        return False
        return has_any_text

    has_text_anywhere = False
    all_cells_bold = True
    for cell in cells:
        if cell.paragraphs and any(p.text.strip() for p in cell.paragraphs):
            has_text_anywhere = True
            if not cell_is_bold(cell):
                all_cells_bold = False
                break
    if has_text_anywhere and all_cells_bold:
        return True
        
    return False

def get_text_format(paragraphs):
    """
    获取主导字体、字号、加粗、对齐等（按众数最频繁出现的值提取）。
    """
    if not paragraphs:
        return {"font": None, "size": None, "bold": None, "align": None}
        
    fonts = []
    sizes = []
    bolds = []
    aligns = []
    
    for p in paragraphs:
        # 1. 字体判定（优先东亚字体 eastAsia）
        p_font = None
        for r in p.runs:
            if r.text.strip():
                r_ea = None
                if hasattr(r, '_element') and r._element is not None:
                    r_fonts = r._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                    if r_fonts is not None:
                        r_ea = r_fonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                p_font = r_ea or r.font.name
                if p_font:
                    break
        if not p_font:
            style = p.style
            for _ in range(10):
                if style is None:
                    break
                style_font = style.font if hasattr(style, 'font') else None
                if style_font:
                    style_ea = None
                    if hasattr(style_font, 'element') and style_font.element is not None:
                        r_fonts = style_font.element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                        if r_fonts is not None:
                            style_ea = r_fonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                    p_font = style_ea or style_font.name
                    if p_font:
                        break
                style = style.base_style if hasattr(style, 'base_style') else None
                
        # 2. 字号判定
        p_size = None
        for r in p.runs:
            if r.text.strip() and r.font.size is not None:
                p_size = f"{r.font.size.pt}pt"
                break
        if not p_size:
            style = p.style
            for _ in range(10):
                if style is None:
                    break
                style_font = style.font if hasattr(style, 'font') else None
                if style_font and style_font.size is not None:
                    p_size = f"{style_font.size.pt}pt"
                    break
                style = style.base_style if hasattr(style, 'base_style') else None
                
        # 3. 加粗判定
        p_bold = None
        for r in p.runs:
            if r.text.strip():
                if r.bold is not None:
                    p_bold = r.bold
                    break
        if p_bold is None:
            style = p.style
            for _ in range(10):
                if style is None:
                    break
                style_font = style.font if hasattr(style, 'font') else None
                if style_font and style_font.bold is not None:
                    p_bold = style_font.bold
                    break
                style = style.base_style if hasattr(style, 'base_style') else None
            if p_bold is None:
                p_bold = False
                
        # 4. 对齐判定
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        p_align_raw = p.paragraph_format.alignment if hasattr(p, 'paragraph_format') and p.paragraph_format else None
        if p_align_raw is None:
            style = p.style
            for _ in range(10):
                if style is None:
                    break
                style_pf = style.paragraph_format if hasattr(style, 'paragraph_format') else None
                if style_pf and style_pf.alignment is not None:
                    p_align_raw = style_pf.alignment
                    break
                style = style.base_style if hasattr(style, 'base_style') else None
        p_align = align_map.get(p_align_raw) if p_align_raw is not None else None
        
        if p_font: fonts.append(p_font)
        if p_size: sizes.append(p_size)
        if p_bold is not None: bolds.append(p_bold)
        if p_align: aligns.append(p_align)
        
    def _most_common(lst, default=None):
        if not lst:
            return default
        return Counter(lst).most_common(1)[0][0]
        
    return {
        "font": _most_common(fonts, None),
        "size": _most_common(sizes, None),
        "bold": _most_common(bolds, False),
        "align": _most_common(aligns, "left")
    }

def _table_fingerprint_key(fp):
    """提取指纹的可哈希唯一表示，用于去重。"""
    def dict_to_tuple(d):
        if d is None:
            return None
        return tuple(sorted((k, v) for k, v in d.items()))
        
    def nested_dict_to_tuple(d):
        if d is None:
            return None
        return tuple(sorted((k, dict_to_tuple(v) if isinstance(v, dict) else v) for k, v in d.items()))

    return (
        fp.get("tbl_style"),
        nested_dict_to_tuple(fp.get("structure")),
        nested_dict_to_tuple(fp.get("border")),
        dict_to_tuple(fp.get("shading")),
        dict_to_tuple(fp.get("cell_margin")),
        dict_to_tuple(fp.get("header_text")),
        dict_to_tuple(fp.get("body_text"))
    )

def extract_table_fingerprints(filepath):
    """
    提取整个 DOCX 的所有表格视觉指纹，去重，并分配 id，以列表返回。
    """
    doc = docx.Document(filepath)
    fingerprints = []
    seen_keys = set()
    
    table_id = 0
    for table in doc.tables:
        tblPr = table._tbl.find(qn('w:tblPr'))
        
        # 1. 结构与表头智能识别
        has_header = detect_has_header_row(table)
        cols = get_table_cols(table)
        rows = len(table.rows)
        
        # 2. 边框提取 (六方向)
        border_info = {
            "top": None,
            "bottom": None,
            "left": None,
            "right": None,
            "insideH": None,
            "insideV": None
        }
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                for border_name in border_info.keys():
                    b_elem = tblBorders.find(qn(f'w:{border_name}'))
                    if b_elem is not None:
                        val = b_elem.get(qn('w:val'))
                        sz = b_elem.get(qn('w:sz'))
                        color = b_elem.get(qn('w:color'))
                        border_info[border_name] = {
                            "val": val,
                            "sz": int(sz) if sz is not None else None,
                            "color": color
                        }
                        
        # 3. 底纹提取
        header_shadings = []
        body_shadings = []
        for idx, row in enumerate(table.rows):
            is_header_row = (idx == 0 and has_header)
            for cell in row.cells:
                shd = get_effective_shading(cell)
                if is_header_row:
                    if shd: header_shadings.append(shd)
                else:
                    if shd: body_shadings.append(shd)
                    
        def _most_common_shd(lst):
            if not lst:
                return None
            return Counter(lst).most_common(1)[0][0]
            
        shading_info = {
            "header": _most_common_shd(header_shadings) if has_header else None,
            "body": _most_common_shd(body_shadings)
        }
        
        # 4. 内边距提取
        cell_margin = {
            "top": "0.0pt",
            "bottom": "0.0pt",
            "left": "5.4pt",
            "right": "5.4pt"
        }
        if tblPr is not None:
            tblCellMar = tblPr.find(qn('w:tblCellMar'))
            if tblCellMar is not None:
                for side in ("top", "bottom", "left", "right"):
                    elem = tblCellMar.find(qn(f'w:{side}'))
                    if elem is not None:
                        w = elem.get(qn('w:w'))
                        if w is not None:
                            try:
                                val_pt = float(w) / 20.0
                                cell_margin[side] = f"{val_pt:.1f}pt"
                            except ValueError:
                                pass
                                
        # 5. 表头与正文文本格式
        header_paras = []
        body_paras = []
        for idx, row in enumerate(table.rows):
            is_header_row = (idx == 0 and has_header)
            for cell in row.cells:
                for p in cell.paragraphs:
                    if is_header_row:
                        header_paras.append(p)
                    else:
                        body_paras.append(p)
                        
        header_text_format = get_text_format(header_paras) if (has_header and header_paras) else None
        body_text_format = get_text_format(body_paras) if body_paras else None
        
        # 6. 表头样本示例
        example_text = ""
        if len(table.rows) > 0:
            example_text = " | ".join(cell.text.strip() for cell in table.rows[0].cells if cell.text.strip())
            if not example_text.strip():
                example_text = " | ".join(cell.text.strip() for cell in table.rows[0].cells)[:100]
                
        # 整合指纹
        fp = {
            "tbl_style": table.style.name,
            "structure": {"cols": cols, "rows": rows, "has_header_row": has_header},
            "border": border_info,
            "shading": shading_info,
            "cell_margin": cell_margin,
            "header_text": header_text_format,
            "body_text": body_text_format,
            "example": example_text
        }
        
        # 7. 去重逻辑
        fp_key = _table_fingerprint_key(fp)
        if fp_key not in seen_keys:
            seen_keys.add(fp_key)
            fp["id"] = table_id
            fingerprints.append(fp)
            table_id += 1
            
    return fingerprints
