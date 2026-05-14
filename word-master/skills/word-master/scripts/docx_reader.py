import docx
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _extract_style_props(style):
    """
    提取 Word 样式的具体属性（字体、字号、对齐、间距等）。
    返回一个包含格式信息的字典。
    """
    props = {}
    if not hasattr(style, 'font'):
        return props
    
    # 提取字体属性
    if style.font.name: props['font'] = style.font.name
    if style.font.size: props['size'] = f"{style.font.size.pt}pt"
    if style.font.bold is not None: props['bold'] = style.font.bold
    if style.font.italic is not None: props['italic'] = style.font.italic
    
    # 提取段落格式属性
    if hasattr(style, 'paragraph_format'):
        pf = style.paragraph_format
        if pf.alignment is not None:
            align_map = {
                WD_ALIGN_PARAGRAPH.LEFT: 'LEFT',
                WD_ALIGN_PARAGRAPH.CENTER: 'CENTER',
                WD_ALIGN_PARAGRAPH.RIGHT: 'RIGHT',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'JUSTIFY'
            }
            props['align'] = align_map.get(pf.alignment, str(pf.alignment))
        if pf.line_spacing is not None: props['line_spacing'] = pf.line_spacing
        if pf.first_line_indent is not None:
            props['first_line_indent'] = f"{pf.first_line_indent.pt}pt"
        if pf.space_before is not None:
            props['space_before'] = f"{pf.space_before.pt}pt"
        if pf.space_after is not None:
            props['space_after'] = f"{pf.space_after.pt}pt"
            
    return props

def _extract_run_overrides(paragraph):
    """
    提取段落中各个 Run（文本块）相对于段落样式的手动格式覆盖。
    """
    overrides = []
    for run in paragraph.runs:
        text = run.text.strip()
        if not text:
            continue
        
        run_props = {}
        # 如果 Run 对象的属性不是 None，说明它被手动显式设置过
        if run.font.name is not None: run_props['font'] = run.font.name
        if run.font.size is not None: run_props['size'] = f"{run.font.size.pt}pt"
        if run.font.bold is not None: run_props['bold'] = run.font.bold
        if run.font.italic is not None: run_props['italic'] = run.font.italic
        
        if run_props:
            prop_str = ", ".join(f"{k}={v}" for k, v in run_props.items())
            overrides.append(f"{prop_str} on \"{text}\"")
    return overrides

def extract_rich_markdown(filepath, overview=False):
    """
    读取 DOCX 文档并将其转换为带有丰富格式信息的 Markdown 格式。
    
    :param filepath: DOCX 文件路径
    :param overview: 如果为 True，则只输出标题大纲（Heading）
    """
    doc = docx.Document(filepath)
    lines = []
    lines.append(f"<!-- DOCUMENT: {filepath} -->")
    
    if not overview:
        # 收集文档中实际使用到的所有样式名称
        used_style_names = set(p.style.name for p in doc.paragraphs if p.text.strip())
        if used_style_names:
            lines.append("<!-- STYLES:")
            for s_name in sorted(used_style_names):
                if s_name in doc.styles:
                    style = doc.styles[s_name]
                    props = _extract_style_props(style)
                    if props:
                        prop_str = ", ".join(f"{k}:{v}" for k, v in props.items())
                        lines.append(f"  {s_name} -> {prop_str}")
                    else:
                        lines.append(f"  {s_name} -> (default)")
                else:
                    lines.append(f"  {s_name} -> (not found in doc.styles)")
            lines.append("-->")
    
    lines.append("") # 头部后的空行
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name
        
        # 在大纲模式下，过滤掉非标题样式的段落
        if overview and not style_name.startswith('Heading'):
            continue
            
        # 提取行内手动格式覆盖（仅在非大纲模式下执行）
        overrides = _extract_run_overrides(para) if not overview else []
        if overrides:
            override_str = " | override: " + "; ".join(overrides)
            lines.append(f"### [p{i}] {text} {{{style_name}{override_str}}}")
        else:
            lines.append(f"### [p{i}] {text} {{{style_name}}}")
        
    return "\n".join(lines)

if __name__ == "__main__":
    filepath = sys.argv[1]
    overview = "--overview" in sys.argv
    print(extract_rich_markdown(filepath, overview))
