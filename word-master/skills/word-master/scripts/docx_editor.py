import docx
import json
import shutil
import sys
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _replace_text_in_runs(paragraph, find_txt, repl_txt):
    """Run-level text replacement to preserve inline formatting (bold, italic, font, etc.)."""
    full_text = paragraph.text
    if find_txt not in full_text:
        return False

    replaced = False
    for run in paragraph.runs:
        if find_txt in run.text:
            run.text = run.text.replace(find_txt, repl_txt)
            replaced = True

    if not replaced:
        paragraph.text = full_text.replace(find_txt, repl_txt)
        replaced = True

    return replaced

def _copy_run_format(src_run, dest_run):
    """Copy formatting from src_run to dest_run"""
    dest_run.bold = src_run.bold
    dest_run.italic = src_run.italic
    dest_run.underline = src_run.underline
    dest_run.style = src_run.style
    if src_run.font:
        if src_run.font.name:
            dest_run.font.name = src_run.font.name
        if src_run.font.size:
            dest_run.font.size = src_run.font.size
        if src_run.font.color and src_run.font.color.rgb:
            dest_run.font.color.rgb = src_run.font.color.rgb
        if src_run.font.highlight_color:
            dest_run.font.highlight_color = src_run.font.highlight_color
        dest_run.font.strike = src_run.font.strike
        dest_run.font.subscript = src_run.font.subscript
        dest_run.font.superscript = src_run.font.superscript

def _apply_font_attributes(run, params):
    if "name" in params:
        run.font.name = params["name"]
    if "east_asia" in params:
        from docx.oxml.ns import qn
        rpr = run._element.find(qn('w:rPr'))
        if rpr is None:
            rpr = run._element.makeelement(qn('w:rPr'), {})
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:eastAsia'), params["east_asia"])
    if "size" in params:
        size_str = str(params["size"])
        if size_str.endswith("pt"):
            from docx.shared import Pt
            run.font.size = Pt(float(size_str.rstrip("pt")))
        elif size_str.isdigit():
            from docx.shared import Pt
            run.font.size = Pt(int(size_str) / 2)
    if "bold" in params:
        run.bold = params["bold"]
    if "italic" in params:
        run.italic = params["italic"]
    if "color_rgb" in params:
        from docx.shared import RGBColor
        hex_str = params["color_rgb"].lstrip("#")
        if hex_str:
            run.font.color.rgb = RGBColor.from_string(hex_str)
        else:
            run.font.color.rgb = None
            from docx.oxml.ns import qn
            rpr = run._element.find(qn('w:rPr'))
            if rpr is not None:
                color_elem = rpr.find(qn('w:color'))
                if color_elem is not None:
                    rpr.remove(color_elem)
    if "color_theme" in params:
        from docx.oxml.ns import qn
        from docx.enum.dml import MSO_THEME_COLOR
        theme_map = {
            "DARK_1": MSO_THEME_COLOR.DARK_1,
            "LIGHT_1": MSO_THEME_COLOR.LIGHT_1,
            "DARK_2": MSO_THEME_COLOR.DARK_2,
            "LIGHT_2": MSO_THEME_COLOR.LIGHT_2,
            "ACCENT_1": MSO_THEME_COLOR.ACCENT_1,
            "ACCENT_2": MSO_THEME_COLOR.ACCENT_2,
            "ACCENT_3": MSO_THEME_COLOR.ACCENT_3,
            "ACCENT_4": MSO_THEME_COLOR.ACCENT_4,
            "ACCENT_5": MSO_THEME_COLOR.ACCENT_5,
            "ACCENT_6": MSO_THEME_COLOR.ACCENT_6,
            "HYPERLINK": MSO_THEME_COLOR.HYPERLINK,
            "FOLLOWED_HYPERLINK": MSO_THEME_COLOR.FOLLOWED_HYPERLINK,
            "TEXT_1": MSO_THEME_COLOR.DARK_1,
            "TEXT_2": MSO_THEME_COLOR.LIGHT_1,
            "BACKGROUND_1": MSO_THEME_COLOR.LIGHT_1,
            "BACKGROUND_2": MSO_THEME_COLOR.LIGHT_2,
        }
        theme_name = params["color_theme"].upper()
        if theme_name in theme_map:
            run.font.color.theme_color = theme_map[theme_name]

def _apply_set_font(paragraph, params):
    """Apply font settings to paragraph runs or specific matched text."""
    text_match = params.get("text_match")
    
    if not text_match:
        for run in paragraph.runs:
            _apply_font_attributes(run, params)
        return

    full_text = paragraph.text
    if text_match not in full_text:
        return
        
    match_index = params.get("match_index", "all")
    match_ranges = []
    start = 0
    while True:
        idx = full_text.find(text_match, start)
        if idx == -1:
            break
        match_ranges.append((idx, idx + len(text_match)))
        start = idx + len(text_match)
        
    if str(match_index) != "all":
        try:
            mi = int(match_index)
            if mi < 0 or mi >= len(match_ranges):
                return
            match_ranges = [match_ranges[mi]]
        except ValueError:
            pass

    if not match_ranges:
        return

    original_runs_info = []
    current_idx = 0
    for r in paragraph.runs:
        r_len = len(r.text)
        original_runs_info.append({
            'text': r.text,
            'start': current_idx,
            'end': current_idx + r_len,
            'run': r
        })
        current_idx += r_len

    paragraph.clear()

    for r_info in original_runs_info:
        r_start = r_info['start']
        r_end = r_info['end']
        r_text = r_info['text']
        
        intersections = []
        for m_start, m_end in match_ranges:
            if m_start < r_end and m_end > r_start:
                intersections.append((max(r_start, m_start), min(r_end, m_end)))
                
        if not intersections:
            new_run = paragraph.add_run(r_text)
            _copy_run_format(r_info['run'], new_run)
            continue
            
        curr_offset = 0
        for i_start, i_end in intersections:
            rel_start = i_start - r_start
            rel_end = i_end - r_start
            
            if rel_start > curr_offset:
                new_run = paragraph.add_run(r_text[curr_offset:rel_start])
                _copy_run_format(r_info['run'], new_run)
                
            new_run = paragraph.add_run(r_text[rel_start:rel_end])
            _copy_run_format(r_info['run'], new_run)
            _apply_font_attributes(new_run, params)
            
            curr_offset = rel_end
            
        if curr_offset < len(r_text):
            new_run = paragraph.add_run(r_text[curr_offset:])
            _copy_run_format(r_info['run'], new_run)

def _apply_set_paragraph_format(paragraph, params):
    """Apply paragraph format settings."""
    pf = paragraph.paragraph_format
    if "alignment" in params:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        pf.alignment = align_map.get(params["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
    if "line_spacing" in params:
        pf.line_spacing = params["line_spacing"]
    if "first_line_indent" in params:
        indent_str = params["first_line_indent"]
        if indent_str.endswith("em"):
            from docx.shared import Emu
            pf.first_line_indent = Emu(int(float(indent_str.rstrip("em")) * 914400 / 2))
        elif indent_str.endswith("pt"):
            from docx.shared import Pt
            pf.first_line_indent = Pt(float(indent_str.rstrip("pt")))
    if "space_before" in params:
        space_str = params["space_before"]
        if space_str.endswith("pt"):
            from docx.shared import Pt
            pf.space_before = Pt(float(space_str.rstrip("pt")))
    if "space_after" in params:
        space_str = params["space_after"]
        if space_str.endswith("pt"):
            from docx.shared import Pt
            pf.space_after = Pt(float(space_str.rstrip("pt")))

def _apply_insert_paragraph(doc, params):
    """Insert a new paragraph before or after the anchor paragraph."""
    content = params.get('content')
    if content is None:
        print("[WARNING] op=insert_paragraph: missing content, skipping.")
        return

    before_id = params.get('before')
    after_id = params.get('after')

    if bool(before_id) == bool(after_id):
        print("[WARNING] op=insert_paragraph: must provide exactly one of 'before' or 'after', skipping.")
        return

    target_id = before_id or after_id
    try:
        idx = int(target_id.replace('p', ''))
        if idx < 0 or idx >= len(doc.paragraphs):
            print(f"[WARNING] op=insert_paragraph: paragraph index {idx} out of bounds, skipping.")
            return
    except (ValueError, AttributeError):
        print("[WARNING] op=insert_paragraph: invalid target format, skipping.")
        return

    anchor_p = doc.paragraphs[idx]
    style_name = params.get('style')

    if before_id:
        new_p = anchor_p.insert_paragraph_before(content)
        if style_name:
            new_p.style = style_name
        elif idx == 0:
            new_p.style = 'Normal'
        else:
            new_p.style = anchor_p.style.name
    else:
        new_p = doc.add_paragraph(content)
        anchor_p._p.addnext(new_p._p)
        new_p.style = style_name if style_name else anchor_p.style.name


def _apply_delete_paragraph(doc, params):
    """Delete the paragraph at the given index."""
    target = params.get('target')
    if not target:
        print("[WARNING] op=delete_paragraph: missing target, skipping.")
        return
    try:
        idx = int(target.replace('p', ''))
        if idx < 0 or idx >= len(doc.paragraphs):
            print(f"[WARNING] op=delete_paragraph: paragraph index {idx} out of bounds, skipping.")
            return
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
    except (ValueError, AttributeError):
        print("[WARNING] op=delete_paragraph: invalid target format, skipping.")


def _apply_modify_cell(doc, params):
    t_idx, r_idx, c_idx = params.get('table'), params.get('row'), params.get('col')
    if any(x is None for x in [t_idx, r_idx, c_idx, params.get('content')]):
        print("[WARNING] op=modify_cell: missing required parameters, skipping.")
        return
    try:
        table = doc.tables[t_idx]
        table.cell(r_idx, c_idx).text = params['content']
    except IndexError:
        print("[WARNING] op=modify_cell: index out of bounds, skipping.")


def _apply_insert_row(doc, params):
    t_idx = params.get('table')
    before, after = params.get('before'), params.get('after')
    if t_idx is None or (before is None and after is None) or (before is not None and after is not None):
        print("[WARNING] op=insert_row: invalid parameters, skipping.")
        return
    try:
        table = doc.tables[t_idx]
        target_row_idx = before if before is not None else after
        target_row = table.rows[target_row_idx]
        new_row = table.add_row()
        if before is not None:
            target_row._tr.addprevious(new_row._tr)
        else:
            target_row._tr.addnext(new_row._tr)
    except IndexError:
        print("[WARNING] op=insert_row: index out of bounds, skipping.")


def _apply_insert_table(doc, params):
    before_id, after_id = params.get('before'), params.get('after')
    rows, cols = params.get('rows'), params.get('cols')
    if (before_id is None and after_id is None) or (before_id and after_id) or rows is None or cols is None:
        print("[WARNING] op=insert_table: invalid parameters, skipping.")
        return
    
    target_id = before_id or after_id
    try:
        idx = int(target_id.replace('p', ''))
        if idx < 0 or idx >= len(doc.paragraphs):
            print(f"[WARNING] op=insert_table: paragraph index {idx} out of bounds, skipping.")
            return
        anchor_p = doc.paragraphs[idx]
        new_table = doc.add_table(rows=rows, cols=cols)
        if before_id:
            anchor_p._p.addprevious(new_table._tbl)
        else:
            anchor_p._p.addnext(new_table._tbl)
    except (ValueError, AttributeError, IndexError):
        print("[WARNING] op=insert_table: invalid target format or out of bounds, skipping.")


def _apply_set_header_footer(doc, params):
    sec_idx = params.get('section', 0)
    try:
        section = doc.sections[sec_idx]
        even = params.get('even_page', False)
        header_text = params.get('header')
        footer_text = params.get('footer')
        
        target_header = section.even_page_header if even else section.header
        target_footer = section.even_page_footer if even else section.footer
        
        if header_text is not None:
            if not target_header.paragraphs:
                target_header.add_paragraph()
            target_header.paragraphs[0].text = header_text
            
        if footer_text is not None:
            if not target_footer.paragraphs:
                target_footer.add_paragraph()
            target_footer.paragraphs[0].text = footer_text
    except IndexError:
        print("[WARNING] op=set_header_footer: section index out of bounds, skipping.")


def _apply_insert_page_break(doc, params):
    before_id, after_id = params.get('before'), params.get('after')
    if (not before_id and not after_id) or (before_id and after_id):
        print("[WARNING] op=insert_page_break: invalid parameters, skipping.")
        return
        
    target_id = before_id or after_id
    try:
        idx = int(target_id.replace('p', ''))
        anchor_p = doc.paragraphs[idx]
        new_p = doc.add_paragraph()
        new_p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        if before_id:
            anchor_p._p.addprevious(new_p._p)
        else:
            anchor_p._p.addnext(new_p._p)
    except (ValueError, AttributeError, IndexError):
        print("[WARNING] op=insert_page_break: invalid target format, skipping.")


def _parse_unit(val_str):
    from docx.shared import Cm, Pt
    if val_str.endswith('cm'):
        return Cm(float(val_str.rstrip('cm')))
    elif val_str.endswith('pt'):
        return Pt(float(val_str.rstrip('pt')))
    return None


def _apply_apply_style(doc, params):
    target = params.get('target')
    style_name = params.get('style')
    if not target or not style_name:
        print("[WARNING] op=apply_style: missing parameters, skipping.")
        return
    try:
        idx = int(target.replace('p', ''))
        if idx < 0 or idx >= len(doc.paragraphs):
            print(f"[WARNING] op=apply_style: paragraph index {idx} out of bounds, skipping.")
            return
        if style_name in doc.styles:
            para = doc.paragraphs[idx]
            para.style = style_name
            if params.get('clear_run_formats'):
                from docx.oxml.ns import qn
                for run in para.runs:
                    run.font.size = None
                    run.bold = None
                    run.italic = None
                    run.font.name = None
                    # Clear both rgb and theme color overrides
                    run.font.color.rgb = None
                    run.font.color.theme_color = None
                    # Clear XML-level overrides (East Asia font, color element)
                    rpr = run._element.find(qn('w:rPr'))
                    if rpr is not None:
                        rfonts = rpr.find(qn('w:rFonts'))
                        if rfonts is not None:
                            rfonts.attrib.pop(qn('w:eastAsia'), None)
                        # Remove w:color element to fully clear color override
                        color_elem = rpr.find(qn('w:color'))
                        if color_elem is not None:
                            rpr.remove(color_elem)
        else:
            print(f"[WARNING] op=apply_style: style '{style_name}' not found, skipping.")
    except (ValueError, AttributeError, IndexError):
        print("[WARNING] op=apply_style: invalid target format or index, skipping.")


def _apply_set_page_setup(doc, params):
    sec_idx = params.get('section', 0)
    try:
        section = doc.sections[sec_idx]
        for margin in ['margin_top', 'margin_bottom', 'margin_left', 'margin_right']:
            if margin in params:
                val = _parse_unit(params[margin])
                if val:
                    if margin == 'margin_top': section.top_margin = val
                    elif margin == 'margin_bottom': section.bottom_margin = val
                    elif margin == 'margin_left': section.left_margin = val
                    elif margin == 'margin_right': section.right_margin = val
        
        orientation = params.get('orientation')
        if orientation == 'landscape':
            section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        elif orientation == 'portrait':
            section.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
    except IndexError:
        print("[WARNING] op=set_page_setup: section index out of bounds, skipping.")


def _apply_update_style_definition(doc, params):
    """Overwrite a built-in style's font/paragraph format with fingerprint values."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    style_name = params.get('style')
    fp = params.get('fingerprint', {})
    if not style_name or not fp:
        print("[WARNING] op=update_style_definition: missing style or fingerprint, skipping.")
        return
    if style_name not in doc.styles:
        print(f"[WARNING] op=update_style_definition: style '{style_name}' not found, skipping.")
        return

    style = doc.styles[style_name]

    # Font size
    size_str = fp.get('size')
    if size_str:
        try:
            style.font.size = Pt(float(size_str.rstrip('pt')))
        except (ValueError, AttributeError):
            pass

    # Bold / italic
    if fp.get('bold') is not None:
        style.font.bold = fp['bold']
    if fp.get('italic') is not None:
        style.font.italic = fp['italic']

    # Western font (name)
    font_name = fp.get('font')
    if font_name:
        style.font.name = font_name
        # Also set East Asia (Chinese) font via XML
        from docx.oxml.ns import qn
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:eastAsia'), font_name)

    # Color
    # python-docx's ColorFormat clears the entire w:color element when setting
    # rgb=None or theme_color=None, so we must manipulate XML directly to set
    # one attribute while clearing the other.
    color = fp.get('color')
    if color:
        from docx.oxml.ns import qn
        rpr = style._element.get_or_add_rPr()
        # Remove existing w:color element to start clean
        for old_color in rpr.findall(qn('w:color')):
            rpr.remove(old_color)
        if color.startswith('rgb:'):
            from docx.shared import RGBColor
            hex_str = color[4:].lstrip('#')
            if hex_str:
                color_elem = rpr.makeelement(qn('w:color'), {
                    qn('w:val'): hex_str,
                })
                rpr.append(color_elem)
        elif color.startswith('theme:'):
            from docx.enum.dml import MSO_THEME_COLOR
            theme_map = {
                "DARK_1": MSO_THEME_COLOR.DARK_1,
                "LIGHT_1": MSO_THEME_COLOR.LIGHT_1,
                "DARK_2": MSO_THEME_COLOR.DARK_2,
                "LIGHT_2": MSO_THEME_COLOR.LIGHT_2,
                "ACCENT_1": MSO_THEME_COLOR.ACCENT_1,
                "ACCENT_2": MSO_THEME_COLOR.ACCENT_2,
                "ACCENT_3": MSO_THEME_COLOR.ACCENT_3,
                "ACCENT_4": MSO_THEME_COLOR.ACCENT_4,
                "ACCENT_5": MSO_THEME_COLOR.ACCENT_5,
                "ACCENT_6": MSO_THEME_COLOR.ACCENT_6,
                "HYPERLINK": MSO_THEME_COLOR.HYPERLINK,
                "FOLLOWED_HYPERLINK": MSO_THEME_COLOR.FOLLOWED_HYPERLINK,
                "TEXT_1": MSO_THEME_COLOR.DARK_1,
                "TEXT_2": MSO_THEME_COLOR.LIGHT_1,
                "BACKGROUND_1": MSO_THEME_COLOR.LIGHT_1,
                "BACKGROUND_2": MSO_THEME_COLOR.LIGHT_2,
            }
            theme_name = color[6:].upper()
            if theme_name in theme_map:
                theme_val = theme_map[theme_name]
                color_elem = rpr.makeelement(qn('w:color'), {
                    qn('w:themeColor'): theme_val.xml_value,
                })
                rpr.append(color_elem)
    elif color == '':
        # Explicitly clear color
        from docx.oxml.ns import qn
        rpr = style._element.find(qn('w:rPr'))
        if rpr is not None:
            for old_color in rpr.findall(qn('w:color')):
                rpr.remove(old_color)

    # Alignment
    align_str = fp.get('align')
    if align_str:
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if align_str in align_map:
            style.paragraph_format.alignment = align_map[align_str]

    # Spacing
    sb = fp.get('space_before')
    if sb:
        try:
            style.paragraph_format.space_before = Pt(float(sb.rstrip('pt')))
        except (ValueError, AttributeError):
            pass

    sa = fp.get('space_after')
    if sa:
        try:
            style.paragraph_format.space_after = Pt(float(sa.rstrip('pt')))
        except (ValueError, AttributeError):
            pass

    # Line spacing (multiplier only)
    ls = fp.get('line_spacing')
    if ls is not None:
        style.paragraph_format.line_spacing = ls


def _backup_file(filepath):
    """Create a .bak backup of the original file before modifying."""
    bak_path = filepath + '.bak'
    shutil.copy2(filepath, bak_path)
    return bak_path

def apply_operations(filepath, ops, outpath=None):
    doc = docx.Document(filepath)
    if outpath is None:
        outpath = filepath

    # Backup original before any modification
    if os.path.exists(filepath):
        _backup_file(filepath)

    for op in ops:
        if op['op'] == 'rewrite_paragraph':
            idx = int(op['target'].replace('p', ''))
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                saved_style = para.style
                para.text = op['content']
                para.style = saved_style
        elif op['op'] == 'replace_text':
            target = op.get('target', 'all')
            find_txt = op['find']
            repl_txt = op['replace']
            if target == 'all':
                for p in doc.paragraphs:
                    _replace_text_in_runs(p, find_txt, repl_txt)
            else:
                idx = int(target.replace('p', ''))
                if idx < len(doc.paragraphs):
                    _replace_text_in_runs(doc.paragraphs[idx], find_txt, repl_txt)
        elif op['op'] == 'set_font':
            idx = int(op['target'].replace('p', ''))
            if idx < len(doc.paragraphs):
                _apply_set_font(doc.paragraphs[idx], op)
        elif op['op'] == 'set_paragraph_format':
            idx = int(op['target'].replace('p', ''))
            if idx < len(doc.paragraphs):
                _apply_set_paragraph_format(doc.paragraphs[idx], op)
        elif op['op'] == 'insert_paragraph':
            _apply_insert_paragraph(doc, op)
        elif op['op'] == 'delete_paragraph':
            _apply_delete_paragraph(doc, op)
        elif op['op'] == 'modify_cell':
            _apply_modify_cell(doc, op)
        elif op['op'] == 'insert_row':
            _apply_insert_row(doc, op)
        elif op['op'] == 'insert_table':
            _apply_insert_table(doc, op)
        elif op['op'] == 'set_header_footer':
            _apply_set_header_footer(doc, op)
        elif op['op'] == 'insert_page_break':
            _apply_insert_page_break(doc, op)
        elif op['op'] == 'apply_style':
            _apply_apply_style(doc, op)
        elif op['op'] == 'set_page_setup':
            _apply_set_page_setup(doc, op)
        elif op['op'] == 'update_style_definition':
            _apply_update_style_definition(doc, op)

    doc.save(outpath)

if __name__ == "__main__":
    filepath = sys.argv[1]
    with open(sys.argv[2], 'r', encoding='utf-8') as f:
        ops = json.load(f)
    outpath = sys.argv[3] if len(sys.argv) > 3 else filepath
    apply_operations(filepath, ops, outpath)
