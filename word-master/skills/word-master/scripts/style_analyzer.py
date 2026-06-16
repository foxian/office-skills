import docx
import json
import re
import sys
import zipfile
from collections import Counter
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CHAPTER_PATTERN = re.compile(
    r'^(第[一二三四五六七八九十\d]+章|[一二三四五六七八九十]+、|（[一二三四五六七八九十\d]+）|\d+\.\s)'
)

WESTERN_FONTS = {
    "times new roman", "arial", "calibri", "cambria", "courier new", 
    "georgia", "helvetica", "segoe ui", "tahoma", "trebuchet ms", 
    "verdana", "symbol", "arial black", "comic sans ms", "lucida sans unicode"
}

def _is_western_font(font_name):
    if not font_name:
        return False
    return font_name.lower().strip() in WESTERN_FONTS

def _resolve_style_attribute(style, attr_name):
    """
    Walk up the style inheritance chain (up to 10 hops) to resolve an attribute
    (e.g., 'font.size', 'font.bold', 'font.italic', 'font.color.rgb', 'font.color.theme_color').
    """
    curr_style = style
    for _ in range(10):
        if curr_style is None:
            break
        parts = attr_name.split('.')
        obj = curr_style
        val = None
        try:
            for part in parts:
                obj = getattr(obj, part) if obj is not None else None
            val = obj
        except AttributeError:
            pass
        if val is not None:
            return val
        curr_style = curr_style.base_style
    return None

def _resolve_style_eastasia_font(style):
    """
    Walk up the style inheritance chain to resolve the eastAsia font.
    """
    curr_style = style
    for _ in range(10):
        if curr_style is None:
            break
        font_obj = curr_style.font if hasattr(curr_style, 'font') else None
        elem = font_obj.element if font_obj and hasattr(font_obj, 'element') else None
        east_asia = _get_eastasia_font(elem)
        if east_asia is not None:
            return east_asia
        curr_style = curr_style.base_style
    return None

def _resolve_paragraph_format_attribute(style, attr_name):
    """
    Walk up the style inheritance chain (up to 10 hops) to resolve a paragraph format attribute
    (e.g., 'space_before', 'space_after', 'first_line_indent', 'line_spacing').
    """
    curr_style = style
    for _ in range(10):
        if curr_style is None:
            break
        pf = curr_style.paragraph_format if hasattr(curr_style, 'paragraph_format') else None
        if pf is not None:
            try:
                val = getattr(pf, attr_name)
                if val is not None:
                    return val
            except AttributeError:
                pass
        curr_style = curr_style.base_style
    return None

_theme_fonts_cache = {}  # filepath -> {major_east_asia, minor_east_asia}


def _load_theme_fonts(filepath):
    """
    Load theme fonts from docx file.
    Returns dict with 'major_east_asia' and 'minor_east_asia' font names.
    """
    if filepath in _theme_fonts_cache:
        return _theme_fonts_cache[filepath]

    result = {"major_east_asia": None, "minor_east_asia": None}
    try:
        with zipfile.ZipFile(filepath, 'r') as z:
            for name in z.namelist():
                if 'theme/theme' in name.lower() and name.endswith('.xml'):
                    content = z.read(name)
                    # 解析 theme XML
                    import xml.etree.ElementTree as ET
                    root = ET.fromstring(content)
                    ns = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}

                    # 查找 majorFont 和 minorFont
                    for font_elem in root.findall('.//a:fontScheme/a:majorFont', ns):
                        ea = font_elem.find('a:ea', ns)
                        if ea is not None and ea.get('typeface'):
                            result['major_east_asia'] = ea.get('typeface')
                        # 查找 script="Hans" (简体中文)
                        for script_font in font_elem.findall('a:font', ns):
                            if script_font.get('script') == 'Hans':
                                result['major_east_asia'] = script_font.get('typeface')
                                break

                    for font_elem in root.findall('.//a:fontScheme/a:minorFont', ns):
                        ea = font_elem.find('a:ea', ns)
                        if ea is not None and ea.get('typeface'):
                            result['minor_east_asia'] = ea.get('typeface')
                        for script_font in font_elem.findall('a:font', ns):
                            if script_font.get('script') == 'Hans':
                                result['minor_east_asia'] = script_font.get('typeface')
                                break
                    break
    except Exception:
        pass

    _theme_fonts_cache[filepath] = result
    return result


def _fingerprint_key(fp):
    """Convert fingerprint dict to a hashable key for grouping."""
    return (fp.get("size"), fp.get("bold"), fp.get("italic"), fp.get("align"), fp.get("font"), fp.get("color"), fp.get("first_line_indent"), fp.get("line_spacing"))


def _get_effective_value(run_val, style_val):
    """Return run value if explicitly set, else fall back to style value."""
    if run_val is not None:
        return run_val
    return style_val


def _get_eastasia_font(element):
    """
    Extract eastAsia font from a font element's XML.
    For Chinese documents, eastAsia font takes priority over ascii font.
    Returns the eastAsia font name, or None if not set.
    """
    if element is None:
        return None
    r_fonts = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if r_fonts is not None:
        east_asia = r_fonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
        if east_asia:
            return east_asia
    return None


def _get_font_name(font_obj):
    """
    Get font name with eastAsia priority for Chinese documents.
    python-docx's font.name only returns ascii font; this checks eastAsia as well.
    """
    if font_obj is None:
        return None, None
    ascii_font = font_obj.name
    elem = font_obj.element if hasattr(font_obj, 'element') else None
    east_asia_font = _get_eastasia_font(elem)
    return ascii_font, east_asia_font


def compute_effective_fingerprint(paragraph, theme_fonts=None):
    """
    Compute effective format fingerprint for a paragraph, merging style attributes and run-level overrides.
    For multi-run paragraphs, uses the first run with an explicit value.
    Returns dict with: size, bold, italic, align, font, color, space_before, space_after, first_line_indent, line_spacing.
    theme_fonts: optional dict with 'major_east_asia' and 'minor_east_asia' from document theme.
    """
    style = paragraph.style
    style_font = style.font if hasattr(style, 'font') else None
    style_pf = style.paragraph_format if hasattr(style, 'paragraph_format') else None

    # Collect run-level overrides (first run with explicit value wins)
    run_size = None
    run_bold = None
    run_italic = None
    run_color_rgb = None
    run_color_theme = None
    for run in paragraph.runs:
        if run.font.size is not None and run_size is None:
            run_size = run.font.size
        if run.bold is not None and run_bold is None:
            run_bold = run.bold
        if run.italic is not None and run_italic is None:
            run_italic = run.italic
        # Color: extract rgb and theme_color
        if run.font.color:
            if run.font.color.rgb is not None and run_color_rgb is None:
                run_color_rgb = str(run.font.color.rgb)
            if run.font.color.theme_color is not None and run_color_theme is None:
                run_color_theme = str(run.font.color.theme_color)

    # Font size (resolved recursively via style inheritance chain)
    style_size = _resolve_style_attribute(style, 'font.size')
    eff_size = run_size or style_size
    size_str = f"{eff_size.pt}pt" if eff_size else None

    # Bold
    style_bold = _resolve_style_attribute(style, 'font.bold')
    eff_bold = _get_effective_value(run_bold, style_bold)

    # Italic
    style_italic = _resolve_style_attribute(style, 'font.italic')
    eff_italic = _get_effective_value(run_italic, style_italic)

    # Alignment
    align_map = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    pf_align = paragraph.paragraph_format.alignment if hasattr(paragraph, 'paragraph_format') and paragraph.paragraph_format else None
    style_align = style_pf.alignment if style_pf else None
    eff_align_raw = _get_effective_value(pf_align, style_align)
    eff_align = align_map.get(eff_align_raw) if eff_align_raw else None

    # Font family: prefer eastAsia font (for Chinese documents), fallback to ascii
    run_ascii = None
    run_east_asia = None
    for run in paragraph.runs:
        if run.font.name is not None and run_ascii is None:
            run_ascii = run.font.name
        run_ea = _get_eastasia_font(run._element)
        if run_ea is not None and run_east_asia is None:
            run_east_asia = run_ea
        if run_ascii is not None and run_east_asia is not None:
            break
    # Resolving font family names through the style inheritance chain
    style_ascii = _resolve_style_attribute(style, 'font.name')
    style_east_asia = _resolve_style_eastasia_font(style)
    # 智能选择主题字体：标题用 major, 正文用 minor
    is_heading = _get_outline_level(paragraph) is not None
    theme_ea = None
    if theme_fonts:
        theme_ea = theme_fonts.get('major_east_asia') if is_heading else theme_fonts.get('minor_east_asia')
        
    # 完美的东亚字体优先回退链：
    # 1. 运行级东亚字体 (例如 run 里硬编码的“黑体”)
    # 2. 样式级东亚字体 (例如 Heading 样式定义的“黑体”)
    # 3. 运行级通用/西文属性，如果它其实是中文字体 (例如 run.font.name="黑体")
    # 4. 样式级通用/西文属性，如果它其实是中文字体
    # 5. 主题级东亚字体 (继承自文档主题，如“等线”)
    # 6. 最后无奈回退到运行级或样式级的西文字体 (如 "Times New Roman")
    if run_east_asia:
        eff_font = run_east_asia
    elif style_east_asia:
        eff_font = style_east_asia
    elif run_ascii and not _is_western_font(run_ascii):
        eff_font = run_ascii
    elif style_ascii and not _is_western_font(style_ascii):
        eff_font = style_ascii
    elif theme_ea:
        eff_font = theme_ea
    elif run_ascii:
        eff_font = run_ascii
    else:
        eff_font = style_ascii

    # Font color: prefer run-level, fall back to style (resolved recursively)
    style_color_rgb = _resolve_style_attribute(style, 'font.color.rgb')
    style_color_theme = _resolve_style_attribute(style, 'font.color.theme_color')
    if style_color_rgb is not None:
        style_color_rgb = str(style_color_rgb)
    if style_color_theme is not None:
        style_color_theme = str(style_color_theme)
    eff_color_rgb = _get_effective_value(run_color_rgb, style_color_rgb)
    eff_color_theme = _get_effective_value(run_color_theme, style_color_theme)
    # Combine into a single color string: "theme:NAME" or "rgb:XXXXXX" or None
    # theme_color has higher priority — it represents semantic intent (e.g., ACCENT_1)
    if eff_color_theme:
        eff_color = f"theme:{eff_color_theme}"
    elif eff_color_rgb:
        eff_color = f"rgb:{eff_color_rgb}"
    else:
        eff_color = None

    # Spacing fields (paragraph-level)

    def _get_pt_str(val):
        """Convert Pt value to 'Xpt' string, defaults to '0.0pt' if not set."""
        if val is None:
            return "0.0pt"
        try:
            return f"{val.pt}pt"
        except AttributeError:
            return "0.0pt"

    sb_val = paragraph.paragraph_format.space_before
    if sb_val is None:
        sb_val = _resolve_paragraph_format_attribute(style, 'space_before')
    eff_space_before = _get_pt_str(sb_val)

    sa_val = paragraph.paragraph_format.space_after
    if sa_val is None:
        sa_val = _resolve_paragraph_format_attribute(style, 'space_after')
    eff_space_after = _get_pt_str(sa_val)

    # First line indent
    fli_val = paragraph.paragraph_format.first_line_indent
    if fli_val is None:
        fli_val = _resolve_paragraph_format_attribute(style, 'first_line_indent')
    eff_first_line_indent = _get_pt_str(fli_val)

    # Line spacing: support both multiplier (float) and absolute (Pt) modes.
    # For absolute/fixed mode, encode as "Xpt/exactly" string so the editor can
    # distinguish it from multiplier mode and apply WD_LINE_SPACING.EXACTLY.
    ls_val = paragraph.paragraph_format.line_spacing
    ls_rule = paragraph.paragraph_format.line_spacing_rule
    # Style-chain line spacing (resolve independently to check for EXACTLY mode)
    style_ls_val = _resolve_paragraph_format_attribute(style, 'line_spacing')
    style_ls_rule = _resolve_paragraph_format_attribute(style, 'line_spacing_rule')
    if ls_val is None:
        ls_val = style_ls_val
        ls_rule = style_ls_rule
    else:
        # If paragraph has a plain multiplier (float) but style chain defines EXACTLY,
        # prefer the style's EXACTLY — paragraph's SINGLE is often a template placeholder reset.
        from docx.enum.text import WD_LINE_SPACING
        para_is_multiplier = not hasattr(ls_val, 'pt')
        style_is_exactly = (style_ls_rule == WD_LINE_SPACING.EXACTLY)
        if para_is_multiplier and style_is_exactly and style_ls_val is not None:
            ls_val = style_ls_val
            ls_rule = style_ls_rule
    eff_line_spacing = None
    if ls_val is not None:
        try:
            pt_val = ls_val.pt  # Pt object → absolute mode
            eff_line_spacing = f"{pt_val:.1f}pt/exactly"
        except AttributeError:
            eff_line_spacing = ls_val  # float → multiplier mode

    return {
        "size": size_str,
        "bold": eff_bold,
        "italic": eff_italic,
        "align": eff_align,
        "font": eff_font,
        "color": eff_color,
        "space_before": eff_space_before,
        "space_after": eff_space_after,
        "first_line_indent": eff_first_line_indent,
        "line_spacing": eff_line_spacing,
    }


def extract_fingerprints(filepath, min_cluster_size=4, heading_aware=False):
    """
    Analyze a DOCX document, extract body format fingerprint clusters.
    Filter out clusters with < min_cluster_size members (cover page, etc).
    Each cluster retains one representative sample text.
    Returns list of {id, fingerprint, example} or {heading_role, fingerprint, example}.

    When heading_aware=True, heading paragraphs are grouped by outline level
    (navigation structure) rather than style name, and their fingerprints are
    averaged across all paragraphs at that level. Body paragraphs still go
    through normal format clustering.
    """
    doc = docx.Document(filepath)
    theme_fonts = _load_theme_fonts(filepath)

    if heading_aware:
        heading_paras = {}    # "Heading N" -> list of fingerprint dicts
        heading_examples = {} # "Heading N" -> first paragraph text
        body_clusters = {}    # fingerprint_key -> list of (text, fp, style_name)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            outline_level = _get_outline_level(para)
            if outline_level is not None:
                role = f"Heading {outline_level + 1}"
                fp = compute_effective_fingerprint(para, theme_fonts)
                if role not in heading_paras:
                    heading_paras[role] = []
                    heading_examples[role] = text
                heading_paras[role].append(fp)
            else:
                # 排除空段落：空白内容不代表任何正文样式，防止封面/页脚空行误主导聚类
                if not text:
                    continue
                fp = compute_effective_fingerprint(para, theme_fonts)
                key = _fingerprint_key(fp)
                if key not in body_clusters:
                    body_clusters[key] = []
                body_clusters[key].append((text, fp, para.style.name if para.style is not None else "Normal"))

        result = []
        for role in sorted(heading_paras, key=lambda r: int(r.split()[-1])):
            result.append({
                "heading_role": role,
                "fingerprint": _average_fingerprints(heading_paras[role]),
                "example": heading_examples[role],
            })
        body_id = 0
        for key, members in body_clusters.items():
            has_chapter_pattern = any(CHAPTER_PATTERN.match(t) for t, _, _ in members)
            if len(members) < min_cluster_size and not has_chapter_pattern:
                continue
            representative_text, representative_fp, _ = members[0]
            
            style_names = [style for _, _, style in members]
            most_common_style = Counter(style_names).most_common(1)[0][0]
            
            result.append({
                "id": body_id,
                "recommended_style": most_common_style,
                "member_count": len(members),
                "fingerprint": representative_fp,
                "example": representative_text
            })
            body_id += 1
        return result

    # --- Original mode (heading_aware=False) ---
    clusters = {}  # key -> list of (paragraph_text, fingerprint, style_name)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        fp = compute_effective_fingerprint(para, theme_fonts)
        key = _fingerprint_key(fp)
        if key not in clusters:
            clusters[key] = []
        clusters[key].append((text, fp, para.style.name if para.style is not None else "Normal"))

    # Strategy A: filter small clusters; Strategy C: chapter pattern clusters exempt from filtering
    result = []
    for i, (key, members) in enumerate(clusters.items()):
        has_chapter_pattern = any(CHAPTER_PATTERN.match(text) for text, _, _ in members)
        if len(members) < min_cluster_size and not has_chapter_pattern:
            continue
        representative_text, representative_fp, _ = members[0]
        
        style_names = [style for _, _, style in members]
        most_common_style = Counter(style_names).most_common(1)[0][0]
        
        result.append({
            "id": i,
            "recommended_style": most_common_style,
            "member_count": len(members),
            "fingerprint": representative_fp,
            "example": representative_text,
        })
    return result


def _get_outline_level(paragraph) -> int | None:
    """
    Read paragraph outline level (0=Heading 1, ..., 8=Heading 9).
    Strategy:
      1. Check paragraph-level XML directly.
      2. Walk style inheritance chain (para.style -> base_style -> ...),
         up to 10 hops to guard against circular references.
    Returns int 0-8 or None (not a heading).
    """
    def _read_outline_lvl_from_pPr(pPr_elem):
        if pPr_elem is None:
            return None
        ol = pPr_elem.find(qn('w:outlineLvl'))
        if ol is None:
            return None
        val = ol.get(qn('w:val'))
        if val is None:
            return None
        try:
            level = int(val)
            return level if 0 <= level <= 8 else None
        except ValueError:
            return None

    # 1. Paragraph-level direct attribute
    para_pPr = paragraph._element.find(qn('w:pPr'))
    level = _read_outline_lvl_from_pPr(para_pPr)
    if level is not None:
        return level

    # 2. Style inheritance chain
    style = paragraph.style
    for _ in range(10):
        if style is None:
            break
        style_pPr = style.element.find(qn('w:pPr'))
        level = _read_outline_lvl_from_pPr(style_pPr)
        if level is not None:
            return level
        style = style.base_style

    return None


def _average_fingerprints(fps: list[dict]) -> dict:
    """
    Aggregate a list of fingerprint dicts into one representative fingerprint.
    Aggregation rules per field: see Spec §3.
    """
    if not fps:
        return {}

    def _pt_to_float(s):
        if s is None:
            return 0.0
        try:
            return float(s.rstrip("pt"))
        except (ValueError, AttributeError):
            return 0.0

    def _vote(values, tiebreak_order=None, skip_none=False):
        filtered = [v for v in values if v is not None] if skip_none else values
        if not filtered:
            return None
        counts = Counter(filtered)
        max_count = max(counts.values())
        winners = [k for k, v in counts.items() if v == max_count]
        if len(winners) == 1:
            return winners[0]
        if tiebreak_order is not None:
            for priority in tiebreak_order:
                if priority in winners:
                    return priority
        # Fallback: first occurrence in input order
        for v in values:
            if v in winners:
                return v
        return winners[0]

    bool_tiebreak = [False, None, True]

    size_vals = [_pt_to_float(fp.get("size")) for fp in fps if fp.get("size") is not None]
    ls_vals = [fp.get("line_spacing") for fp in fps if fp.get("line_spacing") is not None]

    def _avg_pt(field):
        vals = [_pt_to_float(fp.get(field)) for fp in fps]
        return f"{sum(vals) / len(vals):.1f}pt"

    return {
        "size": f"{sum(size_vals) / len(size_vals):.1f}pt" if size_vals else None,
        "bold": _vote([fp.get("bold") for fp in fps], tiebreak_order=bool_tiebreak),
        "italic": _vote([fp.get("italic") for fp in fps], tiebreak_order=bool_tiebreak),
        "align": _vote([fp.get("align") for fp in fps]),
        "font": _vote([fp.get("font") for fp in fps], skip_none=True),
        "color": _vote([fp.get("color") for fp in fps]),
        "space_before": _avg_pt("space_before"),
        "space_after": _avg_pt("space_after"),
        "first_line_indent": _avg_pt("first_line_indent"),
        "line_spacing": _vote(ls_vals, skip_none=True) if ls_vals else None,
    }


def _detect_list_type(paragraph, doc) -> str | None:
    """
    Detect if a paragraph is a Word native list.
    Returns "List Bullet", "List Number", or None.
    Uses <w:numPr> XML node as primary signal; falls back to
    style name keywords and text prefix characters.
    """
    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))

    # 1. 尝试使用 XML 物理节点识别
    if numPr is not None:
        try:
            numId_elem = numPr.find(qn('w:numId'))
            ilvl_elem = numPr.find(qn('w:ilvl'))
            if numId_elem is not None:
                numId = numId_elem.get(qn('w:val'))
                ilvl = ilvl_elem.get(qn('w:val')) if ilvl_elem is not None else "0"

                # 追溯 numbering.xml 定义
                numbering_part = doc.part.numbering_part
                if numbering_part is not None:
                    numbering = numbering_part.numbering_definitions._numbering

                    # 遍历查找 w:num
                    num_node = None
                    for num in numbering.findall(qn('w:num')):
                        if num.get(qn('w:numId')) == numId:
                            num_node = num
                            break

                    if num_node is not None:
                        abstractNumId_elem = num_node.find(qn('w:abstractNumId'))
                        if abstractNumId_elem is not None:
                            abstractNumId = abstractNumId_elem.get(qn('w:val'))

                            # 遍历查找 w:abstractNum
                            abstractNum_node = None
                            for abNum in numbering.findall(qn('w:abstractNum')):
                                if abNum.get(qn('w:abstractNumId')) == abstractNumId:
                                    abstractNum_node = abNum
                                    break

                            if abstractNum_node is not None:
                                # 遍历查找 w:lvl
                                lvl_node = None
                                for lvl in abstractNum_node.findall(qn('w:lvl')):
                                    if lvl.get(qn('w:ilvl')) == ilvl:
                                        lvl_node = lvl
                                        break

                                if lvl_node is not None:
                                    numFmt_elem = lvl_node.find(qn('w:numFmt'))
                                    if numFmt_elem is not None:
                                        numFmt = numFmt_elem.get(qn('w:val'))
                                        if numFmt == 'bullet':
                                            return "List Bullet"
                                        else:
                                            return "List Number"
        except (AttributeError, KeyError, TypeError):
            pass

    # 2. 降级：样式名关键字判定
    style_name = (paragraph.style.name if paragraph.style is not None else "Normal").lower()
    if 'list bullet' in style_name or 'list paragraph' in style_name:
        return "List Bullet"
    if 'list number' in style_name:
        return "List Number"
    if 'list' in style_name:
        if re.match(r'^([0-9a-zA-Z]+|[一二三四五六七八九十]+)[\.、\s]', paragraph.text.strip()):
            return "List Number"
        return "List Bullet"

    # 3. 降级：文本前缀匹配（仅匹配 bullet 符号；不对数字前缀判定以防止拦截标题级联）
    stripped = paragraph.text.strip()
    if stripped.startswith(('•', '·', '-', '*', 'o', '▪', '■', '◆', '▲')):
        return "List Bullet"

    return None


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Extract style fingerprints from a DOCX template.")
    parser.add_argument("filepath", help="Source template DOCX path")
    parser.add_argument("--min-cluster-size", type=int, default=4)
    parser.add_argument("--output", default="fingerprints.json")
    parser.add_argument(
        "--heading-aware",
        action="store_true",
        default=True,
        help=(
            "Group heading paragraphs by outline level (navigation level) instead of style name. "
            "Use when the template has mismatched heading styles but correct navigation structure. "
            "Enabled by default. Use --no-heading-aware to disable."
        ),
    )
    parser.add_argument(
        "--no-heading-aware",
        action="store_false",
        dest="heading_aware",
        help="Disable heading-aware mode (use format clustering instead).",
    )
    args = parser.parse_args()
    fingerprints = extract_fingerprints(
        args.filepath,
        args.min_cluster_size,
        heading_aware=args.heading_aware,
    )
    with open(args.output, "w", encoding="utf-8") as f:
        json.dump(fingerprints, f, ensure_ascii=False, indent=2)
    print(f"[INFO] Extracted {len(fingerprints)} fingerprint groups -> {args.output}")
