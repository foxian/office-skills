import docx
import sys

def convert_markdown(md_path, out_path, template=None):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        else:
            doc.add_paragraph(line)
            
    doc.save(out_path)

if __name__ == "__main__":
    convert_markdown(sys.argv[1], sys.argv[2])
