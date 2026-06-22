
"""
格式转换：Markdown -> DOCX / HTML / TXT / PDF
"""
import argparse
import re
import sys
from pathlib import Path

# 让 print 在 GBK 等非 UTF-8 终端下也能正常输出（特别是 ✓ 字符）
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

sys.path.insert(0, str(Path(__file__).parent))
from _md_utils import read_utf8


def to_html(content):
    import markdown as md_lib
    css = """<style>
body { font-family: 'Segoe UI', Arial, sans-serif; max-width: 720px; margin: 24px auto; padding: 0 16px; line-height: 1.7; }
h1,h2,h3,h4,h5,h6 { border-bottom: 1px solid #eee; padding-bottom: 4px; }
code { background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-family: monospace; }
pre { background: #f4f4f4; padding: 16px; border-radius: 4px; overflow-x: auto; }
blockquote { border-left: 4px solid #ccc; margin: 0; padding: 0 16px; color: #666; }
table { border-collapse: collapse; width: 100%; }
th, td { border: 1px solid #ddd; padding: 8px 12px; }
th { background: #f0f0f0; }
img { max-width: 100%; height: auto; display: block; margin: 16px auto; }
</style>"""
    html_body = md_lib.markdown(content, extensions=["tables", "fenced_code", "toc", "nl2br"])
    return "<!DOCTYPE html>\n<html>\n<head>\n<meta charset='utf-8'>\n" + css + "\n</head>\n<body>\n" + html_body + "\n</body>\n</html>"


def to_txt(content):
    # 移除 frontmatter
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            content = parts[2].strip()
    # 移除 Markdown 标记
    text = re.sub(r"^#{1,6}\s+", "", content, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    text = re.sub(r"\[(.+?)\]\(.*?\)", r"\1", text)
    text = re.sub(r"^[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^>+\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def to_docx(content, input_path, template):
    from docx import Document
    from docx.shared import Inches, Pt
    import markdown as md_lib
    from bs4 import BeautifulSoup
    import io

    if template and Path(template).exists():
        doc = Document(template)
        # 清空正文段落（保留 section/页眉页脚关联以便继承模板版式）
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)
        # 清空正文表格（避免模板里示例表被原样保留）
        for table in list(doc.tables):
            t = table._element
            t.getparent().remove(t)
    else:
        doc = Document()

    base_dir = Path(input_path).parent
    html = md_lib.markdown(content, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")

    for elem in soup.children:
        tag = getattr(elem, "name", None)
        if tag is None:
            continue
        if tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(tag[1])
            doc.add_heading(elem.get_text(), level=level)
        elif tag == "p":
            img = elem.find("img")
            if img:
                src = img.get("src", "")
                if not src.startswith("http"):
                    img_path = (base_dir / src).resolve()
                    if img_path.exists():
                        doc.add_picture(str(img_path), width=Inches(5.5))
                        if img.get("alt"):
                            doc.add_paragraph(img.get("alt"), style="Caption")
            else:
                doc.add_paragraph(elem.get_text())
        elif tag == "ul":
            for li in elem.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif tag == "ol":
            for li in elem.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif tag == "pre":
            code = elem.find("code")
            text = code.get_text() if code else elem.get_text()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif tag == "blockquote":
            doc.add_paragraph(elem.get_text(), style="Quote")
        elif tag == "hr":
            doc.add_paragraph("-" * 40)
        elif tag == "table":
            rows = elem.find_all("tr")
            if not rows:
                continue
            cols = 0
            for row in rows:
                cell_count = len(row.find_all(["td", "th"]))
                if cell_count > cols:
                    cols = cell_count
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for c_idx, cell in enumerate(cells):
                    table.rows[r_idx].cells[c_idx].text = cell.get_text()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _find_browser():
    """按优先级查找可用的浏览器（Edge > Chrome > 环境变量），跨 Windows / macOS / Linux。"""
    import os
    import sys

    if sys.platform == "win32":
        candidates = [
            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
            r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        ]
    elif sys.platform == "darwin":
        candidates = [
            "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
            "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
            "/Applications/Chromium.app/Contents/MacOS/Chromium",
        ]
    else:  # Linux / 其他
        candidates = [
            "/usr/bin/google-chrome",
            "/usr/bin/google-chrome-stable",
            "/usr/bin/chromium",
            "/usr/bin/chromium-browser",
            "/usr/bin/microsoft-edge",
            "/snap/bin/chromium",
        ]

    candidates.append(os.environ.get("CHROME_PATH", ""))
    # 兜底：让系统 PATH 自己找
    candidates.extend(["msedge", "chrome", "chromium", "google-chrome"])

    for c in candidates:
        if not c:
            continue
        if Path(c).is_absolute() and Path(c).exists():
            return c
        # 相对名 / PATH 中的命令：用 shutil.which 探测
        import shutil
        resolved = shutil.which(c)
        if resolved:
            return resolved
    return None


def to_pdf(content, input_path):
    """Markdown → PDF。

    优先使用 weasyprint（纯 Python，跨平台，无外部依赖）；
    若未安装 weasyprint，回退到 Headless Edge/Chrome。
    """
    html_content = to_html(content)
    # 注入 <base> 标签，让 HTML 中的相对路径（图片等）以 markdown 源文件目录为基准解析
    base_href = Path(input_path).resolve().parent.as_uri() + "/"
    html_content = html_content.replace(
        "<head>",
        f"<head>\n<base href=\"{base_href}\">",
        1,
    )

    # 方案 1：weasyprint（纯 Python）
    try:
        from weasyprint import HTML
        return HTML(string=html_content, base_url=base_href).write_pdf()
    except ImportError:
        pass  # 未安装，回退到浏览器
    except Exception as e:
        # weasyprint 存在但渲染失败（例如系统缺少某些字体/库），给出提示并回退
        print(f"[warn] weasyprint 渲染失败 ({e})，回退到 Headless 浏览器。", file=sys.stderr)

    # 方案 2：Headless Edge/Chrome（需要系统已安装浏览器）
    import subprocess
    import tempfile

    browser = _find_browser()
    if not browser:
        raise RuntimeError(
            "未找到可用的 PDF 生成方案。请任选其一：\n"
            "  1. 安装纯 Python 方案：  pip install weasyprint  （推荐，跨平台）\n"
            "  2. 安装 Microsoft Edge 或 Google Chrome 后重试\n"
            "  3. 将浏览器可执行文件所在目录加入系统 PATH\n"
            "  4. 设置环境变量 CHROME_PATH 指向浏览器可执行文件"
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_html = Path(tmpdir) / "input.html"
        tmp_pdf = Path(tmpdir) / "output.pdf"
        tmp_html.write_text(html_content, encoding="utf-8")

        cmd = [
            str(browser),
            "--headless=new",
            "--disable-gpu",
            "--no-sandbox",
            "--no-pdf-header-footer",  # 关闭页脚 URL 水印
            f"--print-to-pdf={tmp_pdf}",
            f"file://{tmp_html.resolve()}",
        ]
        try:
            proc = subprocess.run(cmd, check=True, timeout=120, capture_output=True, text=True)
        except FileNotFoundError as e:
            raise RuntimeError(f"浏览器不可执行: {browser}") from e
        except subprocess.CalledProcessError as e:
            raise RuntimeError(
                f"浏览器调用失败 (exit={e.returncode}):\nstdout: {e.stdout}\nstderr: {e.stderr}"
            ) from e

        if not tmp_pdf.exists():
            raise RuntimeError("浏览器未生成 PDF 文件")
        return tmp_pdf.read_bytes()


def convert_file(input_path, fmt, output_path, template):
    content = read_utf8(str(input_path))

    if fmt == "html":
        result = to_html(content)
        output_path.write_text(result, encoding="utf-8")
    elif fmt == "txt":
        result = to_txt(content)
        output_path.write_text(result, encoding="utf-8")
    elif fmt == "docx":
        result = to_docx(content, str(input_path), template)
        output_path.write_bytes(result)
    elif fmt == "pdf":
        result = to_pdf(content, str(input_path))
        output_path.write_bytes(result)

    print("✓ " + str(input_path) + " -> " + str(output_path))


def main():
    parser = argparse.ArgumentParser(description="Markdown 格式转换工具")
    parser.add_argument("input", help="输入 Markdown 文件或目录")
    parser.add_argument("--to", required=True, choices=["docx", "html", "txt", "pdf"], dest="fmt", help="目标格式")
    parser.add_argument("-o", "--output", help="输出路径（文件或目录）")
    parser.add_argument("--template", help="DOCX 模板文件路径")

    args = parser.parse_args()
    input_path = Path(args.input)

    if input_path.is_dir():
        md_files = list(input_path.rglob("*.md"))
        out_dir = Path(args.output) if args.output else input_path / (args.fmt + "_output")
        out_dir.mkdir(parents=True, exist_ok=True)
        for md_file in md_files:
            rel = md_file.relative_to(input_path)
            out_file = out_dir / rel.with_suffix("." + args.fmt)
            out_file.parent.mkdir(parents=True, exist_ok=True)
            convert_file(md_file, args.fmt, out_file, args.template)
    else:
        if args.output:
            out_file = Path(args.output)
        else:
            out_file = input_path.with_suffix("." + args.fmt)
        out_file.parent.mkdir(parents=True, exist_ok=True)
        convert_file(input_path, args.fmt, out_file, args.template)


if __name__ == "__main__":
    main()
